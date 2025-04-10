# single_file_app.py

import asyncio
import base64
import concurrent.futures
import json
import logging
import os
import re
import tempfile
import threading
import queue
import atexit
import time
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import fitz  # PyMuPDF
import numpy as np
import pandas as pd
import pdfplumber  # For PDF viewer rendering - Can likely be removed if fitz rendering is stable
import streamlit as st
import torch  # Usually implicitly required by sentence-transformers
from docx import Document as DocxDocument  # Renamed to avoid conflict
from docx import Document # Import for Word export
from docx.shared import Pt, RGBColor # Import for Word export styling
from docx.enum.text import WD_ALIGN_PARAGRAPH # Import for Word export alignment
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, util
from thefuzz import fuzz
import openai  # Added for Azure OpenAI integration
import zipfile
import urllib.parse
import io
import sys
import math
import uuid
import torch
import docx
import logging
import asyncio
import threading
import traceback
import numpy as np
import pandas as pd
import tempfile
import concurrent.futures
from PIL import Image
from copy import deepcopy
from typing import List, Dict, Any, Tuple, Optional, Union, Set
from pathlib import Path
from datetime import datetime, timedelta
import streamlit as st
from streamlit import session_state as ss
import fitz  # PyMuPDF
from sentence_transformers import SentenceTransformer
import openai  # Added for Azure OpenAI integration
import google.generativeai as genai  # Added for Google Gemini integration
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# --- Configuration ---
logging.basicConfig(
    level=logging.INFO, # Changed default level to INFO, DEBUG for detailed tracing
    format="%(asctime)s - %(threadName)s - %(name)s - %(levelname)s - %(message)s" # Added threadName
)
logger = logging.getLogger(__name__)
load_dotenv()

# ****** SET PAGE CONFIG HERE (First Streamlit command) ******
st.set_page_config(layout="wide", page_title="SmartDocs")
# ****** END SET PAGE CONFIG ******

# --- Load images as base64 for embedding in CSS ---
def get_base64_encoded_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()

# Load all the required images
try:
    mascot_base64 = get_base64_encoded_image("assets/mascot.png")
    # Remove ifcontrollers logo
    images_loaded = True
    logger.info("Successfully loaded all brand images")
except Exception as e:
    logger.error(f"Failed to load one or more images: {e}")
    images_loaded = False
    mascot_base64 = ""

# Create checkmark icon for features list
check_base64 = get_base64_encoded_image("assets/correct.png")
check_img = f'<img src="data:image/png;base64,{check_base64}" class="check-icon" alt="✓">' if check_base64 else "✅"

# --- Custom CSS Styling ---
# Define primary colors
PROCESS_CYAN = "#00ADE4"
DARK_BLUE = "#002345"
LIGHT_BLUE_TINT = "#E6F7FD" # Example tint (adjust as needed)
VERY_LIGHT_GRAY = "#FAFAFA"

st.markdown(f"""
<style>
    /* Base Styling */
    .stApp {{
        background-color: white; /* Light background */
    }}

    /* Headings */
    h1, h2, h3, h4, h5, h6 {{
        color: {DARK_BLUE};
    }}
    h1 {{ /* Center Main Title */
        text-align: center;
    }}

    /* Buttons */
    .stButton > button[kind="primary"] {{
        background-color: {PROCESS_CYAN};
        color: white;
        border: 1px solid {PROCESS_CYAN};
    }}
    .stButton > button[kind="primary"]:hover {{
        background-color: {DARK_BLUE};
        color: white;
        border: 1px solid {DARK_BLUE};
    }}
     .stButton > button[kind="primary"]:disabled {{
        background-color: #cccccc;
        color: #666666;
        border: 1px solid #cccccc;
    }}

    .stButton > button[kind="secondary"] {{
        color: {DARK_BLUE};
        border: 1px solid {DARK_BLUE};
    }}
    .stButton > button[kind="secondary"]:hover {{
        border-color: {PROCESS_CYAN};
        color: {PROCESS_CYAN};
        background-color: rgba(0, 173, 228, 0.1); /* Light Process Cyan background */
    }}
     .stButton > button[kind="secondary"]:disabled {{
        color: #aaaaaa;
        border-color: #dddddd;
    }}

    /* Expanders */
    .stExpander > summary {{
        background-color: {LIGHT_BLUE_TINT};
        color: {DARK_BLUE};
        border-radius: 0.25rem;
        border: 1px solid rgba(0, 173, 228, 0.2); /* Subtle border */
    }}
    .stExpander > summary:hover {{
        background-color: rgba(0, 173, 228, 0.2); /* Slightly darker tint on hover */
    }}
    .stExpander > summary svg {{ /* Expander Icon Color */
        fill: {DARK_BLUE};
    }}

    /* Container Borders */
    /* Targeting containers used for analysis sections and chat */
    .st-emotion-cache-1r6slb0 {{ /* This selector might be fragile, adjust if needed */
        border: 1px solid {LIGHT_BLUE_TINT};
    }}
    .st-emotion-cache-lrl5gf {{ /* This selector might be fragile, adjust if needed */
         border: 1px solid {LIGHT_BLUE_TINT};
    }}

    /* Download Button */
    .stDownloadButton > button {{
        background-color: {DARK_BLUE};
        color: white;
        border: 1px solid {DARK_BLUE};
    }}
    .stDownloadButton > button:hover {{
        background-color: {PROCESS_CYAN};
        color: {DARK_BLUE};
        border: 1px solid {PROCESS_CYAN};
    }}

    /* Text Input / Area */
    .stTextInput, .stTextArea {{
        border-color: rgba(0, 35, 69, 0.2); /* Subtle dark blue border */
    }}

    /* Mascot Image - Fixed Position at Bottom Right */
    .mascot-image {{
        position: fixed;
        bottom: 20px;
        right: 20px;
        width: 150px; /* Adjust size as needed */
        z-index: 0; /* Lower z-index to ensure it's in the background */
        opacity: 0.7; /* Slightly more transparent */
        pointer-events: none; /* Prevents the mascot from intercepting mouse clicks */
    }}

    
    .smartdocs-subtitle {{
        color: {DARK_BLUE};
        font-size: 1.2rem;
        font-weight: 400;
        text-align: center;
        margin-top: 0.3rem;
        margin-bottom: 2rem;
    }}

    /* Hide default title */
    .st-emotion-cache-10trblm {{
        display: none;
    }}
    
    /* Features Container Styling */
    .features-container {{
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        width: 100%;
        padding: 20px;
        border-radius: 10px;
    }}
    
    .features-row {{
        display: flex;
        justify-content: space-between;
        width: 80%;
        margin: 10px 0;
    }}
    
    .feature-text {{
        width: 48%;
        background-color: white;
        padding: 15px;
        border-radius: 8px;
        box-shadow: 2px 2px 10px rgba(0, 0, 0, 0.1);
        font-size: 14px;
        font-family: Arial, sans-serif;
        display: flex;
        align-items: center;
        border: 1px solid {LIGHT_BLUE_TINT};
    }}
    
    .welcome-header {{
        color: {DARK_BLUE};
        font-size: 24px;
        text-align: center;
        margin-bottom: 20px;
        font-weight: 500;
    }}
</style>

<!-- Fixed Mascot Image -->
<div class="mascot-image">
    <img src="data:image/png;base64,{mascot_base64}" alt="Mascot">
</div>

<!-- SmartDocs Title (replacing logo) -->
<div>
    
</div>
""", unsafe_allow_html=True)

st.markdown(
    "<h1 style='text-align: center; color: #00ADE4; margin-left: 1rem;'>SmartDocs</h1>",
    unsafe_allow_html=True,
)
# --- End Custom CSS Styling ---


MAX_WORKERS = int(os.environ.get("MAX_WORKERS", 4))
ENABLE_PARALLEL = os.environ.get("ENABLE_PARALLEL", "true").lower() == "true"
FUZZY_MATCH_THRESHOLD = 88  # Adjust this threshold (0-100)
RAG_TOP_K = 10 # Number of relevant chunks to retrieve per sub-prompt (Adjusted from 15)
LOCAL_EMBEDDING_MODEL_PATH = "./embedding_model_local"
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"  # Smaller, faster model for embeddings

# Using Google Gemini instead of Azure OpenAI
DECOMPOSITION_MODEL_NAME = "gemini-2.0-flash-Lite"  # Using gemini-2.0-flash for decomposition
ANALYSIS_MODEL_NAME = "gemini-2.0-flash"  # Using gemini-2.0-flash for analysis

# --- End Google Gemini configuration ---

# --- Load Embedding Model (Cached) ---
@st.cache_resource # Use cache_resource for non-data objects like models
def load_embedding_model():
    """Loads the SentenceTransformer model and caches it."""
    model = None
    if not os.path.exists(LOCAL_EMBEDDING_MODEL_PATH) or not os.path.isdir(LOCAL_EMBEDDING_MODEL_PATH):
        logger.error(f"Local embedding model directory not found at: {LOCAL_EMBEDDING_MODEL_PATH}")
        st.error(f"Fatal Error: Embedding model not found at the required local path ({LOCAL_EMBEDDING_MODEL_PATH}). Please ensure the model directory is present.")
        # Depending on requirements, you might raise an Exception or return None carefully
        return None # Stop the app if the local model is essential

    try:
        logger.info(f"Loading embedding model: {EMBEDDING_MODEL_NAME}...")
        # Check for CUDA availability, fallback to CPU if needed
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model = SentenceTransformer(LOCAL_EMBEDDING_MODEL_PATH)
        logger.info("Embedding model loaded successfully from local path.")
    except Exception as e:
        logger.error(f"Error loading embedding model from {LOCAL_EMBEDDING_MODEL_PATH}: {e}", exc_info=True)
        st.error(f"Failed to load the embedding model from the local path. Error: {e}")
        model = None # Ensure model is None on error
    return model

# Load the model using the cached function
embedding_model = load_embedding_model()

# Function to preprocess files when uploaded
def preprocess_file(file_data: bytes, filename: str, use_advanced_extraction: bool = False):
    """
    Preprocesses a file by extracting chunks and computing their embeddings.
    Stores the results in session state for later use during prompt processing.
    
    Args:
        file_data: Raw bytes of the uploaded file
        filename: Name of the file
        use_advanced_extraction: Optional flag for advanced extraction features
        
    Returns:
        dict: Dictionary with preprocessing status and message
    """
    if embedding_model is None:
        logger.error(f"Skipping preprocessing for {filename}: Embedding model not loaded.")
        return {"status": "error", "message": "Embedding model not loaded"}
        
    try:
        logger.info(f"Starting preprocessing for {filename}")
        file_extension = Path(filename).suffix.lower()
        
        # Extract chunks based on file type
        if file_extension == ".pdf":
            processor = PDFProcessor(file_data)
            chunks, full_text = processor.extract_structured_text_and_chunks()
            original_pdf_bytes = file_data
        elif file_extension == ".docx":
            word_processor = WordProcessor(file_data)
            pdf_bytes = word_processor.convert_to_pdf_bytes()
            if not pdf_bytes: 
                raise ValueError("Failed to convert DOCX to PDF.")
            processor = PDFProcessor(pdf_bytes)
            chunks, full_text = processor.extract_structured_text_and_chunks()
            original_pdf_bytes = pdf_bytes
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
        if not chunks:
            logger.warning(f"No chunks extracted for {filename} during preprocessing.")
            return {"status": "warning", "message": "No text chunks could be extracted"}
            
        # Generate embeddings for all chunks
        logger.info(f"Generating embeddings for {len(chunks)} chunks from {filename}")
        chunk_texts = [chunk.get("text", "") for chunk in chunks]
        valid_chunk_indices = [i for i, text in enumerate(chunk_texts) if text.strip()]
        valid_chunk_texts = [chunk_texts[i] for i in valid_chunk_indices]
        
        if not valid_chunk_texts:
            logger.warning(f"No valid chunk texts found for {filename} during preprocessing.")
            return {"status": "warning", "message": "No valid chunk texts found"}
            
        chunk_embeddings = embedding_model.encode(
            valid_chunk_texts, convert_to_tensor=True, show_progress_bar=False
        )
        
        # Store preprocessed data in session state
        if "preprocessed_data" not in st.session_state:
            st.session_state.preprocessed_data = {}
            
        st.session_state.preprocessed_data[filename] = {
            "chunks": chunks,
            "chunk_embeddings": chunk_embeddings,
            "valid_chunk_indices": valid_chunk_indices,
            "original_bytes": original_pdf_bytes,
            "timestamp": datetime.now().isoformat()
        }
        
        logger.info(f"Successfully preprocessed {filename} with {len(chunks)} chunks and embeddings.")
        return {"status": "success", "message": f"Preprocessed {len(chunks)} chunks"}
        
    except Exception as e:
        logger.error(f"Error preprocessing {filename}: {str(e)}", exc_info=True)
        return {"status": "error", "message": f"Preprocessing failed: {str(e)}"}


# --- Helper Functions ---

def normalize_text(text: Optional[str]) -> str:
    """Normalize text for comparison: lowercase, strip, whitespace."""
    if not text:
        return ""
    text = str(text)
    text = text.lower()  # Case-insensitive matching
    text = re.sub(r"\s+", " ", text)  # Normalize whitespace
    return text.strip()

def remove_markdown_formatting(text: Optional[str]) -> str:
    """Removes common markdown formatting."""
    if not text:
        return ""
    text = str(text)
    # Basic bold, italics, code
    text = re.sub(r"\*(\*|_)(.*?)\1\*?", r"\2", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    # Basic headings, blockquotes, lists
    text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\>\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[\*\-\+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    return text.strip()

def get_base64_encoded_image(image_path: str) -> Optional[str]:
    """Get base64 encoded image."""
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode()
    except Exception as e:
        logger.error(f"Error encoding image {image_path}: {str(e)}")
        return None

def run_async(coro):
    """Helper function to run async code in a thread-safe manner."""
    try:
        loop = asyncio.get_event_loop_policy().get_event_loop()
        if loop.is_running():
            # If a loop is running, create a future and run in executor
            with concurrent.futures.ThreadPoolExecutor() as pool:
                future = pool.submit(asyncio.run, coro)
                return future.result()
        else:
            return loop.run_until_complete(coro)
    except RuntimeError:
        logger.info("No current event loop found, creating a new one.")
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        return loop.run_until_complete(coro)
    except Exception as e:
         logger.error(f"Error running async task: {e}", exc_info=True)
         raise

# Thread-safe counter class (Optional, can be removed if not used)
class Counter:
    def __init__(self, initial_value=0):
        self._value = initial_value
        self._lock = threading.Lock()

    def increment(self):
        with self._lock:
            self._value += 1
            return self._value

    def value(self):
        with self._lock:
            return self._value

# --- RAG Retrieval Function ---
def retrieve_relevant_chunks(
    prompt: str,
    chunks: List[Dict[str, Any]],
    model: SentenceTransformer,
    top_k: int,
    precomputed_embeddings=None,
    valid_chunk_indices=None
) -> List[Dict[str, Any]]:
    """
    Retrieves the top_k most relevant chunks based on semantic similarity to the prompt.
    If precomputed embeddings are provided, skips chunk embedding generation.
    
    Args:
        prompt: The prompt/query text to find matches for
        chunks: List of document chunks
        model: The SentenceTransformer model for embeddings
        top_k: Number of top chunks to retrieve
        precomputed_embeddings: Optional precomputed embeddings for chunks
        valid_chunk_indices: Optional mapping of valid chunk indices
        
    Returns:
        List[Dict[str, Any]]: List of dictionaries, each containing 'text', 'page_num', 'score', 'chunk_id' for the relevant chunks.
    """
    if not chunks or not prompt or model is None:
        logger.warning(f"RAG retrieval skipped for prompt '{prompt[:50]}...': No chunks, prompt, or model.")
        return []

    chunk_texts = [chunk.get("text", "") for chunk in chunks]
    if not any(chunk_texts):
        logger.warning(f"RAG retrieval skipped for prompt '{prompt[:50]}...': All chunk texts are empty.")
        return []

    try:
        # Generate embedding for the prompt
        logger.info(f"RAG: Generating embedding for prompt '{prompt[:50]}...'")
        prompt_embedding = model.encode(
            prompt, convert_to_tensor=True, show_progress_bar=False
        )
        
        # Use precomputed embeddings if available, otherwise compute them now
        if precomputed_embeddings is not None and valid_chunk_indices is not None:
            logger.info(f"RAG: Using precomputed embeddings for {len(precomputed_embeddings)} chunks")
            chunk_embeddings = precomputed_embeddings
        else:
            logger.info(f"RAG: Generating embeddings for {len(chunk_texts)} chunks")
            valid_chunk_indices = [i for i, text in enumerate(chunk_texts) if text.strip()]
            valid_chunk_texts = [chunk_texts[i] for i in valid_chunk_indices]

            if not valid_chunk_texts:
                logger.warning(f"RAG retrieval skipped for prompt '{prompt[:50]}...': No non-empty chunk texts.")
                return []

            chunk_embeddings = model.encode(
                valid_chunk_texts, convert_to_tensor=True, show_progress_bar=False
            )

        # Ensure embeddings are on the same device
        if prompt_embedding.device != chunk_embeddings.device:
            prompt_embedding = prompt_embedding.to(chunk_embeddings.device)
            logger.debug(f"Moved prompt embedding to device: {chunk_embeddings.device}")

        cosine_scores = util.pytorch_cos_sim(prompt_embedding, chunk_embeddings)[0]
        cosine_scores_np = cosine_scores.cpu().numpy()

        actual_top_k = min(top_k, len(valid_chunk_indices))
        # Avoid error if actual_top_k is 0
        if actual_top_k == 0:
             logger.warning(f"RAG: No valid chunks to retrieve for prompt '{prompt[:50]}...'.")
             return []
        top_k_indices_relative = np.argpartition(cosine_scores_np, -actual_top_k)[-actual_top_k:]
        top_k_scores = cosine_scores_np[top_k_indices_relative]
        sorted_top_k_indices_relative = top_k_indices_relative[np.argsort(top_k_scores)[::-1]]

        top_k_original_indices = [
            valid_chunk_indices[i] for i in sorted_top_k_indices_relative
        ]

        # --- Construct list of result dictionaries ---
        results = []
        for i, chunk_index in enumerate(top_k_original_indices):
            chunk = chunks[chunk_index]
            score = cosine_scores_np[sorted_top_k_indices_relative[i]]
            results.append({
                "text": chunk.get("text", ""),
                "page_num": chunk.get("page_num", -1), # Use -1 or None if page unknown
                "score": float(score), # Convert numpy float
                "chunk_id": chunk.get("chunk_id", f"unknown_{chunk_index}")
            })

        logger.info(
            f"RAG: Retrieved {len(results)} chunks for prompt '{prompt[:50]}...'."
        )

        # --- Return the list of dictionaries ---
        return results

    except Exception as e:
        logger.error(f"Error during RAG retrieval for prompt '{prompt[:50]}...': {e}", exc_info=True)
        return []


# --- NEW: Multi-Document RAG Retrieval for Chat ---
def retrieve_relevant_chunks_for_chat(
    prompt: str,
    top_k_per_doc: int,
) -> List[Dict[str, Any]]:
    """
    Retrieves the top_k most relevant chunks from ALL processed documents
    based on semantic similarity to the chat prompt.

    Args:
        prompt: The chat prompt/query text.
        top_k_per_doc: Number of top chunks to retrieve *per document*.

    Returns:
        List[Dict[str, Any]]: List of dictionaries, each containing 'filename',
                              'text', 'page_num', 'score', 'chunk_id' for the
                              most relevant chunks across all documents.
    """
    global embedding_model # Access the globally loaded model
    if embedding_model is None:
        logger.error("Chat RAG skipped: Embedding model not loaded.")
        return []

    if "preprocessed_data" not in st.session_state or not st.session_state.preprocessed_data:
        logger.warning("Chat RAG skipped: No preprocessed documents found in session state.")
        return []

    all_relevant_chunks = []
    logger.info(f"Starting chat RAG for prompt '{prompt[:50]}...' across {len(st.session_state.preprocessed_data)} documents.")

    for filename, data in st.session_state.preprocessed_data.items():
        if not data or 'chunks' not in data or 'chunk_embeddings' not in data:
            logger.warning(f"Skipping document {filename} for chat RAG: Missing required preprocessed data.")
            continue

        logger.debug(f"Running RAG for chat prompt on {filename}...")
        try:
            # Use the refactored retrieve_relevant_chunks for this document
            doc_relevant_chunks = retrieve_relevant_chunks(
                prompt=prompt,
                chunks=data.get("chunks", []),
                model=embedding_model,
                top_k=top_k_per_doc,
                precomputed_embeddings=data.get("chunk_embeddings"),
                valid_chunk_indices=data.get("valid_chunk_indices")
            )

            # Add filename to each retrieved chunk
            for chunk in doc_relevant_chunks:
                chunk['filename'] = filename
                all_relevant_chunks.append(chunk)

            logger.debug(f"Retrieved {len(doc_relevant_chunks)} relevant chunks from {filename} for chat.")

        except Exception as e:
            logger.error(f"Error retrieving chunks for chat from {filename}: {e}", exc_info=True)

    # Optional: Sort all combined chunks by score (highest first)
    all_relevant_chunks.sort(key=lambda x: x.get('score', 0), reverse=True)

    # Optional: Limit the total number of chunks sent to the LLM
    # TOTAL_CHAT_CONTEXT_LIMIT = 20 # Example limit
    # all_relevant_chunks = all_relevant_chunks[:TOTAL_CHAT_CONTEXT_LIMIT]

    logger.info(f"Chat RAG finished. Found {len(all_relevant_chunks)} potentially relevant chunks across all documents.")
    return all_relevant_chunks


# --- AI Analyzer Class ---
_thread_local = threading.local()

class DocumentAnalyzer:
    def __init__(self):
        pass  # Lazy initialization

    def _ensure_client(self, model_name: str):
        """Ensure that the AI client is initialized for the current thread for the specific model."""
        # Store clients per model name in thread local storage
        if not hasattr(_thread_local, "clients"):
            _thread_local.clients = {}

        if model_name not in _thread_local.clients:
            try:
                # --- Google Gemini Configuration ---
                if hasattr(st, "secrets"):
                    # Get Google API key from secrets
                    google_api_key = st.secrets.get("GOOGLE_API_KEY")
                    logger.info(f"Using Google Gemini configuration from Streamlit secrets for model {model_name}.")
                else:
                    # Get Google API key from environment variables
                    google_api_key = os.getenv("GOOGLE_API_KEY")
                    
                    if not google_api_key:
                        raise ValueError("Google Gemini API key is missing.")
                    
                    logger.info(f"Using Google Gemini configuration from environment variables for model {model_name}.")
                
                # Configure the Gemini client
                genai.configure(api_key=google_api_key)
                
                # No need to store client as genai is configured globally
                _thread_local.clients[model_name] = {
                    "client": genai,
                    "model_name": model_name
                }
                
                logger.info(
                    f"Initialized Google Gemini client for thread {threading.current_thread().name} "
                    f"with model: {model_name}"
                )
                
            except Exception as e:
                logger.error(f"Error initializing AI client for model {model_name}: {str(e)}")
                raise
        return _thread_local.clients[model_name]

    async def _get_completion(
        self, messages: List[Dict[str, str]], model_name: str
    ) -> str:
        """Helper method to get completion from the specified AI model."""
        try:
            client_data = self._ensure_client(model_name)
            client = client_data["client"]
            
            # --- Google Gemini API Call ---
            # Convert to Gemini message format
            gemini_messages = []
            for msg in messages:
                role = msg.get("role")
                content = msg.get("content")
                
                # Gemini uses 'user' and 'model' roles
                if role == "system":
                    # Prepend system messages to first user message or add as user message if no user message exists
                    system_content = content
                    if any(m.get("role") == "user" for m in messages):
                        continue  # We'll handle this when processing user messages
                    else:
                        gemini_messages.append({"role": "user", "parts": [{"text": system_content}]})
                elif role == "user":
                    # Check if we need to prepend a system message
                    user_content = content
                    system_content = ""
                    if not gemini_messages and any(m.get("role") == "system" for m in messages):
                        system_msg = next((m for m in messages if m.get("role") == "system"), None)
                        if system_msg:
                            system_content = system_msg.get("content", "")
                    
                    text_content = f"{system_content}\n\n{user_content}" if system_content else user_content
                    gemini_messages.append({"role": "user", "parts": [{"text": text_content}]})
                elif role == "assistant" or role == "model":
                    gemini_messages.append({"role": "model", "parts": [{"text": content}]})
            
            logger.info(f"Sending request to Google Gemini model: {model_name}")
            
            # Create Gemini model
            model = client.GenerativeModel(model_name)
            
            # Use asynchronous event loop with a future to run the synchronous Gemini API call
            loop = asyncio.get_event_loop()
            response_future = loop.run_in_executor(
                None, 
                lambda: model.generate_content(
                    gemini_messages,
                    generation_config=client.GenerationConfig(
                        temperature=0.1,  # Keep low for factual tasks
                        max_output_tokens=8192,  # Max for analysis
                    )
                )
            )
            
            # Wait for the future to complete
            response = await response_future
            
            if not response or not hasattr(response, "text"):
                logger.error(f"Google Gemini ({model_name}) returned no text.")
                raise ValueError(f"Google Gemini ({model_name}) returned no text.")
            
            content = response.text
            logger.info(f"Received response from Google Gemini model {model_name}")
            return content
            
        except Exception as e:
            logger.error(
                f"Error getting completion from AI model {model_name}: {str(e)}", exc_info=True
            )
            if hasattr(e, "response") and hasattr(e.response, "text"):
                logger.error(f"API Error Response Text: {e.response.text}")
            # Specific check for timeout errors (common with complex tasks)
            if "Timeout" in str(e) or "DeadlineExceeded" in str(e):
                 raise TimeoutError(f"API request timed out for model {model_name}.") from e
            raise

    async def decompose_prompt(self, user_prompt: str) -> List[Dict[str, str]]: # Changed return type
        """
        Analyzes the user prompt with an LLM to break it down into individual questions/tasks,
        each with a suggested concise title.
        Returns a list of dictionaries, each containing 'sub_prompt' and 'title'.
        Returns [{'sub_prompt': user_prompt, 'title': 'Overall Analysis'}] on failure.
        """
        logger.info(f"Decomposing prompt: '{user_prompt[:100]}...'")
        system_prompt = """You are a helpful assistant. Your task is to analyze the user's prompt and identify distinct questions or analysis tasks within it.
Break down the prompt into a list of self-contained, individual questions or tasks. For each task, also provide a concise, descriptive title (max 5-6 words).
Your entire response MUST be a single JSON object containing a single key "decomposition", whose value is a list of JSON objects. Each object in the list must have two keys: "title" (the concise title) and "sub_prompt" (the full sub-prompt text).
Do not include any explanations, introductory text, or markdown formatting outside the JSON structure.

Example Input Prompt:
"what was the change in the median net housing value?
was there any change in The homeownership rate?
how many families had credit card debt?
how many families had student debt?"

Example JSON Output:
{
  "decomposition": [
    {
      "title": "Median Net Housing Value Change",
      "sub_prompt": "what was the change in the median net housing value?"
    },
    {
      "title": "Homeownership Rate Change",
      "sub_prompt": "was there any change in The homeownership rate?"
    },
    {
      "title": "Families with Credit Card Debt",
      "sub_prompt": "how many families had credit card debt?"
    },
    {
      "title": "Families with Student Debt",
      "sub_prompt": "how many families had student debt?"
    }
  ]
}

Example Input Prompt:
"Analyze the termination clause and liability limitations."

Example JSON Output:
{
  "decomposition": [
    {
      "title": "Termination Clause Analysis",
      "sub_prompt": "Analyze the termination clause."
    },
    {
      "title": "Liability Limitations Analysis",
      "sub_prompt": "Analyze the liability limitations."
    }
  ]
}

Example Input Prompt:
"Summarize the key findings of the report."

Example JSON Output:
{
  "decomposition": [
    {
      "title": "Key Findings Summary",
      "sub_prompt": "Summarize the key findings of the report."
    }
  ]
}
"""
        human_prompt = f"Analyze the following prompt and return the decomposed questions/tasks and their titles as a JSON list of objects according to the system instructions:\n\n{user_prompt}"

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": human_prompt},
        ]

        # Fallback result in case of errors
        fallback_result = [{"title": "Overall Analysis", "sub_prompt": user_prompt}]

        try:
            response_content = await self._get_completion(messages, model_name=DECOMPOSITION_MODEL_NAME)

            # Attempt to parse the JSON
            try:
                # Clean potential markdown fences (same logic as before)
                cleaned_response = response_content.strip()
                match = re.search(r"```json\s*(\{.*?\})\s*```", cleaned_response, re.DOTALL)
                if match:
                    json_str = match.group(1)
                elif cleaned_response.startswith("{") and cleaned_response.endswith("}"):
                    json_str = cleaned_response
                else:
                     first_brace = cleaned_response.find('{')
                     last_brace = cleaned_response.rfind('}')
                     if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                         json_str = cleaned_response[first_brace:last_brace+1]
                         logger.warning("Used basic brace finding for JSON extraction in decomposition.")
                     else:
                         raise json.JSONDecodeError("Could not find JSON structure.", cleaned_response, 0)

                parsed_json = json.loads(json_str)

                # Validate the new structure
                if isinstance(parsed_json, dict) and "decomposition" in parsed_json:
                    decomposition_list = parsed_json["decomposition"]
                    if (isinstance(decomposition_list, list) and
                        all(isinstance(item, dict) and "title" in item and "sub_prompt" in item
                            and isinstance(item["title"], str) and isinstance(item["sub_prompt"], str)
                            for item in decomposition_list)):

                        logger.info(f"Successfully decomposed prompt into {len(decomposition_list)} sub-prompts with titles.")
                        # Filter out items with empty titles or sub-prompts
                        valid_items = [item for item in decomposition_list if item["title"].strip() and item["sub_prompt"].strip()]
                        if not valid_items:
                             logger.warning("Decomposition resulted in an empty list after filtering. Falling back.")
                             return fallback_result
                        return valid_items
                    else:
                        logger.warning("Decomposition JSON found, but 'decomposition' key is not a list of valid {'title': str, 'sub_prompt': str} objects. Falling back.")
                        return fallback_result
                else:
                    logger.warning("Decomposition JSON parsed, but missing 'decomposition' key or wrong structure. Falling back.")
                    return fallback_result

            except json.JSONDecodeError as json_err:
                logger.error(f"Failed to parse decomposition response as JSON: {json_err}. Raw response: {response_content}")
                logger.warning("Falling back to using the original prompt due to JSON parsing error.")
                return fallback_result

        except TimeoutError:
            logger.error(f"Prompt decomposition request timed out. Falling back to original prompt.")
            return fallback_result
        except Exception as e:
            logger.error(f"Error during prompt decomposition LLM call: {str(e)}", exc_info=True)
            logger.warning("Falling back to using the original prompt.")
            return fallback_result


    @property
    def output_schema_analysis(self) -> dict:
        """Defines the expected JSON structure for document analysis."""
        # Keep this schema definition as it's used in the analysis prompt
        return {
            "title": "Concise Title for the Analysis Section based on the specific sub-prompt",
            "analysis_sections": {
                "descriptive_section_name_1": {
                    "Analysis": "Detailed analysis text for this section...",
                    "Supporting_Phrases": [
                        "Exact quote 1 from the document text...",
                        "Exact quote 2, potentially longer...",
                    ],
                    "Context": "Optional context about this section (e.g., source sub-prompt)",
                },
                # Add more sections as identified by the AI FOR THIS SUB-PROMPT
            },
        }

    async def analyze_document(
        self, relevant_document_text: str, filename: str, sub_prompt: str # Renamed parameter
    ) -> str:
        """
        Analyzes the *relevant document excerpts* based on a specific *sub-prompt* using RAG.
        Returns a JSON string containing the analysis for *this specific sub-prompt*.
        """
        try:
            if not relevant_document_text:
                logger.warning(
                    f"Skipping AI analysis for sub-prompt '{sub_prompt[:50]}...' in {filename}: No relevant text provided."
                )
                # Return a structured message indicating skipped analysis for this sub-prompt
                return json.dumps(
                    {
                        "sub_prompt_analyzed": sub_prompt,
                        "analysis_summary": f"Analysis skipped for sub-prompt '{sub_prompt}' because no relevant text sections were identified.",
                        "supporting_quotes": ["No relevant phrase found."],
                        "analysis_context": "RAG Retrieval Found No Matches"
                    },
                    indent=2,
                )

            # New simplified schema directly in the prompt
            simplified_schema = {
                "sub_prompt_analyzed": "The exact sub-prompt being analyzed",
                "analysis_summary": "Detailed analysis directly answering the sub-prompt...",
                "supporting_quotes": [
                    "Exact quote 1 from the document text...",
                    "Exact quote 2, potentially longer..."
                ],
                "analysis_context": "Optional context about the analysis (e.g., document section names)"
            }

            schema_str = json.dumps(simplified_schema, indent=2)

            # ***** MODIFIED SYSTEM PROMPT FOR RAG & SUB-PROMPT *****
            system_prompt = f"""You are an intelligent document analyser specializing in legal and financial documents. You will be given **relevant excerpts** from a document, identified based on a SPECIFIC USER SUB-PROMPT. Your task is to analyze ONLY these provided excerpts based ONLY on the given SUB-PROMPT and provide a SINGLE, FOCUSED analysis following a specific JSON schema.

### Core Instructions:
1.  **Focus on the Sub-Prompt:** Your entire analysis must address *only* the specific user SUB-PROMPT provided below. Ignore information in the excerpts not relevant to this particular sub-prompt.
2.  **Analyze Thoroughly:** Read the sub-prompt and the document excerpts carefully. Perform the requested analysis based *only* on the excerpts.
3.  **Strict JSON Output:** Your entire response MUST be a single JSON object matching the schema provided below. Do not include any introductory text, explanations, apologies, or markdown formatting (` ```json`, ` ``` `) outside the JSON structure.
4.  **Direct Answer:** Provide a *single, comprehensive analysis* that directly answers the sub-prompt. DO NOT break down your response into multiple sections or categories.
5.  **Exact Supporting Quotes:** The `supporting_quotes` array must contain *only direct, verbatim quotes* from the 'Relevant Document Excerpts'. Preserve original case, punctuation, and formatting. Do *not* include excerpt metadata (Chunk ID/Page/Score) in the quotes. Aim for complete sentences or meaningful clauses.
6.  **No Quote Found:** If no relevant phrase *within the provided excerpts* directly supports your analysis, include the exact string "No relevant phrase found." in the `supporting_quotes` array.
7.  **Focus *Only* on Excerpts:** Base your analysis *exclusively* on the text provided under '### Relevant Document Excerpts:'. Do not infer information not present in these specific excerpts.
8.  **Context Field:** Optionally use the `analysis_context` field to note the relevant section of the document.

### JSON Output Schema:
```json
{schema_str}
```
"""

            human_prompt = f"""Please analyze the following document based *only* on the user sub-prompt below, using ONLY the relevant excerpts provided.

Document Name:
{filename}

User Sub-Prompt:
{sub_prompt}

Relevant Document Excerpts (Identified for the sub-prompt above):
--- START EXCERPTS ---
{relevant_document_text}
--- END EXCERPTS ---

Generate a SINGLE, FOCUSED analysis that directly answers the sub-prompt, strictly following the JSON schema provided in the system instructions. Ensure the analysis *only* addresses this specific sub-prompt and supporting quotes are exact quotes from the TEXT portion of the excerpts provided above."""

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": human_prompt},
            ]

            logger.info(f"Sending RAG analysis request for sub-prompt '{sub_prompt[:50]}...' in {filename} to AI ({ANALYSIS_MODEL_NAME}).")
            # logger.debug(f"Relevant text being sent to AI for {filename} (sub-prompt '{sub_prompt[:50]}...'):\n---\n{relevant_document_text}\n---")

            response_content = await self._get_completion(messages, model_name=ANALYSIS_MODEL_NAME)
            logger.info(f"Received AI analysis response for sub-prompt '{sub_prompt[:50]}...' in {filename}.")

            # --- JSON Parsing and Cleaning ---
            try:
                cleaned_response = response_content.strip()
                match = re.search(r"```json\s*(\{.*?\})\s*```", cleaned_response, re.DOTALL)
                if match:
                    json_str = match.group(1)
                elif cleaned_response.startswith("{") and cleaned_response.endswith("}"):
                    json_str = cleaned_response
                else:
                    first_brace = cleaned_response.find('{')
                    last_brace = cleaned_response.rfind('}')
                    if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
                        json_str = cleaned_response[first_brace:last_brace+1]
                        logger.warning("Used basic brace finding for JSON extraction in analysis.")
                    else:
                         raise json.JSONDecodeError("Could not find JSON structure in analysis response.", cleaned_response, 0)

                parsed_json = json.loads(json_str)

                # Basic Schema Validation
                if not isinstance(parsed_json, dict):
                    raise ValueError("AI response is not a valid JSON object.")
                
                # Check for required fields in new schema
                required_fields = ["sub_prompt_analyzed", "analysis_summary", "supporting_quotes"]
                missing_fields = [field for field in required_fields if field not in parsed_json]
                
                if missing_fields:
                    logger.error(f"AI analysis response missing required fields: {missing_fields}")
                    # Attempt to salvage if possible
                    for field in missing_fields:
                        if field == "sub_prompt_analyzed":
                            parsed_json["sub_prompt_analyzed"] = sub_prompt
                        elif field == "analysis_summary":
                            parsed_json["analysis_summary"] = "Failed to generate analysis for this sub-prompt."
                        elif field == "supporting_quotes":
                            parsed_json["supporting_quotes"] = ["No relevant phrase found."]
                    
                    logger.warning(f"Salvaged analysis response by adding missing fields: {missing_fields}")
                
                # Ensure supporting_quotes is a list
                if not isinstance(parsed_json.get("supporting_quotes", []), list):
                    parsed_json["supporting_quotes"] = ["No relevant phrase found."]
                    logger.warning("Converted non-list supporting_quotes to default list.")

                logger.info(f"Successfully parsed AI analysis JSON for sub-prompt '{sub_prompt[:50]}...'.")
                return json.dumps(parsed_json, indent=2)

            except json.JSONDecodeError as json_err:
                logger.error(f"Failed to parse AI analysis response for sub-prompt '{sub_prompt[:50]}...' as JSON: {json_err}")
                logger.error(f"Raw response content was:\n{response_content}")
                # Return error JSON
                error_response = {
                    "sub_prompt_analyzed": sub_prompt,
                    "analysis_summary": f"Failed to parse AI response as JSON. Error: {json_err}. See logs for raw response.",
                    "supporting_quotes": ["No relevant phrase found."],
                    "analysis_context": "JSON Parsing Error"
                }
                return json.dumps(error_response, indent=2)
            except ValueError as val_err:
                logger.error(f"Error validating AI analysis response structure for sub-prompt '{sub_prompt[:50]}...': {val_err}")
                error_response = {
                    "sub_prompt_analyzed": sub_prompt,
                    "analysis_summary": f"AI response structure validation failed: {val_err}",
                    "supporting_quotes": ["No relevant phrase found."],
                    "analysis_context": "Validation Error"
                }
                return json.dumps(error_response, indent=2)

        except TimeoutError:
             logger.error(f"AI analysis request timed out for sub-prompt '{sub_prompt[:50]}...' in {filename}.")
             error_response = {
                "sub_prompt_analyzed": sub_prompt,
                "analysis_summary": f"The analysis request for this sub-prompt timed out.",
                "supporting_quotes": ["No relevant phrase found."],
                "analysis_context": "Request Timeout"
             }
             return json.dumps(error_response, indent=2)

        except Exception as e:
            logger.error(
                f"Error during AI document analysis for sub-prompt '{sub_prompt[:50]}...' in {filename}: {str(e)}",
                exc_info=True,
            )
            error_response = {
                "sub_prompt_analyzed": sub_prompt,
                "analysis_summary": f"An unexpected error occurred during analysis: {str(e)}",
                "supporting_quotes": ["No relevant phrase found."],
                "analysis_context": "System Error"
            }
            return json.dumps(error_response, indent=2)


    # --- NEW: Chat Response Generation --- 
    async def generate_chat_response(
        self, user_prompt: str, relevant_chunks: List[Dict[str, Any]]
    ) -> str:
        """
        Generates a conversational response to a user query based on relevant chunks
        retrieved from multiple documents using RAG.

        Args:
            user_prompt: The user's chat message.
            relevant_chunks: A list of dictionaries, each containing chunk details
                             ('filename', 'text', 'page_num', 'score').

        Returns:
            A string containing the AI's conversational response, potentially including
            citations like (Source: filename.pdf, Page: 5).
        """
        if not relevant_chunks:
            logger.warning(f"No relevant context found for chat prompt: '{user_prompt[:50]}...'. Returning default response.")
            return "I couldn't find specific information related to your question in the provided documents. Could you try rephrasing or asking something else?"

        # --- Format Context for LLM --- 
        context_str = "\n\n---\n\n".join(
            f"Source: {chunk.get('filename', 'Unknown')}, Page: {chunk.get('page_num', -1) + 1}\n"
            f"Score: {chunk.get('score', 0):.3f}\n"
            f"Content: {chunk.get('text', '')}"
            for chunk in relevant_chunks
        )

        # --- Define System Prompt for Chat --- 
        system_prompt = f"""You are a helpful AI assistant designed to answer questions about documents. You will be given a user's question and relevant excerpts from one or more documents.

Your task is:
1. Understand the user's question.
2. Analyze the provided document excerpts to find the answer.
3. Generate a clear, concise, and conversational response based *only* on the information in the excerpts.
4. **Crucially, you MUST cite your sources.** When you use information from an excerpt, add an inline citation immediately after the information, formatted *exactly* as: `(Source: [filename], Page: [page_number])`. Use the filename and page number provided with each excerpt.
5. If multiple excerpts support a statement, you can list multiple citations like `(Source: doc1.pdf, Page: 2)(Source: doc2.pdf, Page: 5)`.
6. If the provided excerpts do not contain the answer to the user's question, explicitly state that you couldn't find the information in the provided context.
7. Do not make assumptions or provide information not present in the excerpts.
8. Respond directly to the user's question without preamble like "Based on the context...".

Example Excerpt Format:
Source: contract_A.pdf, Page: 5
Score: 0.850
Content: The termination clause allows for a 30-day notice period.

Example Response:
The contract allows for a 30-day notice period (Source: contract_A.pdf, Page: 5)."""

        # --- Construct Messages for LLM --- 
        human_prompt = f"""User Question: {user_prompt}

Relevant Document Excerpts:
---
{context_str}
---

Please answer the user's question based *only* on these excerpts and cite your sources accurately using the format (Source: [filename], Page: [page_number])."""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": human_prompt},
        ]

        try:
            logger.info(f"Sending chat request for prompt '{user_prompt[:50]}...' to AI ({ANALYSIS_MODEL_NAME}). Context length: {len(context_str)} chars.")
            response_content = await self._get_completion(messages, model_name=ANALYSIS_MODEL_NAME)
            logger.info(f"Received chat response for prompt '{user_prompt[:50]}...'.")
            # Basic cleaning (optional)
            response_content = response_content.strip()
            return response_content

        except TimeoutError:
            logger.error(f"Chat request timed out for prompt '{user_prompt[:50]}...'.")
            return "I apologize, but the request timed out while generating a response. Please try again."
        except Exception as e:
            logger.error(f"Error during chat AI call for prompt '{user_prompt[:50]}...': {str(e)}", exc_info=True)
            return f"Sorry, I encountered an error while trying to generate a response: {str(e)}"


# --- Document Processors ---
class PDFProcessor:
    """Handles PDF processing, chunking, verification, and annotation."""

    def __init__(self, pdf_bytes: bytes):
        if not isinstance(pdf_bytes, bytes):
            raise ValueError("pdf_bytes must be of type bytes")
        self.pdf_bytes = pdf_bytes
        self._chunks: List[Dict[str, Any]] = []
        self._full_text: Optional[str] = None
        self._processed = False # Flag to track if extraction ran
        logger.info(f"PDFProcessor initialized with {len(pdf_bytes)} bytes.")

    @property
    def chunks(self) -> List[Dict[str, Any]]:
        if not self._processed:
            self.extract_structured_text_and_chunks()  # Lazy extraction
        return self._chunks

    # Keep full_text property in case it's needed elsewhere
    @property
    def full_text(self) -> str:
        if not self._processed:
            self.extract_structured_text_and_chunks()  # Lazy extraction
        return self._full_text if self._full_text is not None else ""

    def extract_structured_text_and_chunks(self) -> Tuple[List[Dict[str, Any]], str]:
        """Extracts text using PyMuPDF blocks and groups them into chunks."""
        if self._processed:  # Already processed
             return self._chunks, self._full_text if self._full_text is not None else ""

        self._chunks = []
        all_text_parts = []
        current_chunk_id = 0
        doc = None
        try:
            logger.info("Starting structured text extraction and chunking...")
            doc = fitz.open(stream=self.pdf_bytes, filetype="pdf")
            logger.info(f"PDF opened with {doc.page_count} pages.")

            min_chunk_char_length = 50 # Chunks must be at least this long

            for page_num, page in enumerate(doc):
                try:
                    blocks = page.get_text("blocks")
                    blocks.sort(key=lambda b: (b[1], b[0]))

                    current_chunk_text_parts = []
                    current_chunk_bboxes = []
                    last_y1 = 0

                    for b in blocks:
                        x0, y0, x1, y1, text, block_no, block_type = b
                        block_text = text.strip()
                        if not block_text:
                            continue

                        block_rect = fitz.Rect(x0, y0, x1, y1)
                        vertical_gap = y0 - last_y1 if current_chunk_text_parts else 0
                        # Simple paragraph logic: new chunk if big vertical gap
                        is_new_paragraph = not current_chunk_text_parts or (vertical_gap > 8) # Threshold

                        if is_new_paragraph and current_chunk_text_parts:
                            # Finalize previous chunk
                            chunk_text_content = " ".join(current_chunk_text_parts).strip()
                            if len(normalize_text(chunk_text_content)) >= min_chunk_char_length:
                                chunk_id = f"chunk_{current_chunk_id}"
                                self._chunks.append({
                                    "chunk_id": chunk_id,
                                    "text": chunk_text_content,
                                    "page_num": page_num,
                                    "bboxes": current_chunk_bboxes,
                                })
                                all_text_parts.append(chunk_text_content)
                                current_chunk_id += 1
                            # Start new chunk
                            current_chunk_text_parts = [block_text]
                            current_chunk_bboxes = [block_rect]
                        else:
                            # Add to current chunk
                            current_chunk_text_parts.append(block_text)
                            current_chunk_bboxes.append(block_rect)

                        last_y1 = y1

                    # Save the last chunk of the page if it meets length requirement
                    if current_chunk_text_parts:
                        chunk_text_content = " ".join(current_chunk_text_parts).strip()
                        if len(normalize_text(chunk_text_content)) >= min_chunk_char_length:
                            chunk_id = f"chunk_{current_chunk_id}"
                            self._chunks.append({
                                "chunk_id": chunk_id,
                                "text": chunk_text_content,
                                "page_num": page_num,
                                "bboxes": current_chunk_bboxes,
                            })
                            all_text_parts.append(chunk_text_content)
                            current_chunk_id += 1

                except Exception as page_err:
                    logger.error(f"Error processing page {page_num}: {page_err}")

            self._full_text = "\n\n".join(all_text_parts)
            self._processed = True
            logger.info(
                f"Extraction complete. Generated {len(self._chunks)} chunks. "
                f"Total text length: {len(self._full_text or '')} chars."
            )

        except Exception as e:
            logger.error(f"Failed to extract text/chunks: {str(e)}", exc_info=True)
            self._full_text = ""
            self._chunks = []
            self._processed = True # Mark as processed even on failure

        finally:
            if doc:
                doc.close()
        return self._chunks, self._full_text if self._full_text is not None else ""


    def verify_and_locate_phrases(
        self, ai_analysis_json_str: str # Expects the *aggregated* JSON string
    ) -> Tuple[Dict[str, bool], Dict[str, List[Dict[str, Any]]]]:
        """Verifies AI phrases from the aggregated analysis against chunks and locates them."""
        verification_results = {}
        phrase_locations = {}

        chunks_data = self.chunks
        if not chunks_data:
            logger.warning("No chunks available for verification.")
            return {}, {}

        try:
            # Parse the *aggregated* AI analysis
            ai_analysis = json.loads(ai_analysis_json_str)

            # Check if the entire analysis was just an error placeholder
            if not ai_analysis.get("analysis_sections") or \
               all(k.startswith("error_") for k in ai_analysis.get("analysis_sections", {})):
                logger.warning("AI analysis contains only errors or is empty, skipping phrase verification.")
                return {}, {}

            phrases_to_verify = set()
            # Extract all supporting phrases from *all* sections in the aggregated analysis
            for section_key, section_data in ai_analysis.get("analysis_sections", {}).items():
                 # Skip sections indicating skipped RAG or errors generated during analysis
                 if section_key.startswith("info_skipped_") or section_key.startswith("error_"):
                     continue
                 if isinstance(section_data, dict):
                    # Check for both old and new field names for supporting phrases
                    phrases = section_data.get("Supporting_Phrases", section_data.get("supporting_quotes", []))
                    if isinstance(phrases, list):
                        for phrase in phrases:
                            p_text = ""
                            if isinstance(phrase, str):
                                p_text = phrase
                            p_text = p_text.strip()
                            # Exclude the "No relevant phrase found." placeholder
                            if p_text and p_text != "No relevant phrase found.":
                                phrases_to_verify.add(p_text)

            if not phrases_to_verify:
                logger.info("No supporting phrases found in aggregated AI analysis to verify.")
                return {}, {}

            logger.info(
                f"Starting verification for {len(phrases_to_verify)} unique phrases "
                f"(from aggregated analysis) against {len(chunks_data)} original chunks."
            )

            normalized_chunks = [
                (chunk, normalize_text(chunk["text"])) for chunk in chunks_data if chunk.get("text")
            ]

            doc = None
            try:
                doc = fitz.open(stream=self.pdf_bytes, filetype="pdf")

                for original_phrase in phrases_to_verify:
                    verification_results[original_phrase] = False
                    phrase_locations[original_phrase] = []
                    normalized_phrase = normalize_text(remove_markdown_formatting(original_phrase))
                    if not normalized_phrase: continue

                    found_match_for_phrase = False
                    best_score_for_phrase = 0

                    # Verify against ALL original chunks
                    for chunk, norm_chunk_text in normalized_chunks:
                        if not norm_chunk_text: continue

                        score = fuzz.partial_ratio(normalized_phrase, norm_chunk_text)
                        # score = fuzz.token_set_ratio(normalized_phrase, norm_chunk_text) # Alternative

                        if score >= FUZZY_MATCH_THRESHOLD:
                            if not found_match_for_phrase:
                                logger.info(f"Verified (Score: {score}) '{original_phrase[:60]}...' potentially in chunk {chunk['chunk_id']}")
                            found_match_for_phrase = True
                            verification_results[original_phrase] = True
                            best_score_for_phrase = max(best_score_for_phrase, score)

                            # --- Precise Location Search ---
                            page_num = chunk["page_num"]
                            if 0 <= page_num < doc.page_count:
                                page = doc[page_num]
                                clip_rect = fitz.Rect()
                                for bbox in chunk.get('bboxes', []):
                                    try:
                                        if isinstance(bbox, fitz.Rect): clip_rect.include_rect(bbox)
                                        elif isinstance(bbox, (list, tuple)) and len(bbox) == 4: clip_rect.include_rect(fitz.Rect(bbox))
                                    except Exception as bbox_err: logger.warning(f"Skipping invalid bbox {bbox} in chunk {chunk['chunk_id']}: {bbox_err}")

                                if not clip_rect.is_empty:
                                    try:
                                        cleaned_search_phrase = remove_markdown_formatting(original_phrase)
                                        cleaned_search_phrase = re.sub(r"\s+", " ", cleaned_search_phrase).strip()
                                        instances = page.search_for(cleaned_search_phrase, clip=clip_rect, quads=False)

                                        if instances:
                                            logger.debug(f"Found {len(instances)} instance(s) via search_for in chunk {chunk['chunk_id']} area for '{cleaned_search_phrase[:60]}...'")
                                            for rect in instances:
                                                if isinstance(rect, fitz.Rect) and not rect.is_empty:
                                                    phrase_locations[original_phrase].append({
                                                        "page_num": page_num,
                                                        "rect": [rect.x0, rect.y0, rect.x1, rect.y1],
                                                        "chunk_id": chunk["chunk_id"],
                                                        "match_score": score,
                                                        "method": "exact_cleaned_search",
                                                    })
                                        else:
                                            # Fallback to chunk bounding box if exact search fails within the verified chunk
                                            logger.debug(f"Exact search failed for '{cleaned_search_phrase[:60]}...' in verified chunk {chunk['chunk_id']} (score: {score}). Falling back to chunk bbox.")
                                            phrase_locations[original_phrase].append({
                                                "page_num": page_num,
                                                "rect": [clip_rect.x0, clip_rect.y0, clip_rect.x1, clip_rect.y1],
                                                "chunk_id": chunk["chunk_id"],
                                                "match_score": score,
                                                "method": "fuzzy_chunk_fallback",
                                            })
                                    except Exception as search_err: logger.error(f"Error during search_for/fallback in chunk {chunk['chunk_id']}: {search_err}")
                            # else: logger.warning(f"Invalid page number {page_num} for chunk {chunk['chunk_id']}")

                    if not found_match_for_phrase:
                        logger.warning(f"NOT Verified: '{original_phrase[:60]}...' did not meet fuzzy threshold ({FUZZY_MATCH_THRESHOLD}) in any chunk.")
            finally:
                if doc: doc.close()

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse aggregated AI analysis JSON for verification: {e}")
            logger.debug(f"Problematic JSON string: {ai_analysis_json_str[:500]}...") # Log start of bad JSON
        except Exception as e:
            logger.error(f"Error during phrase verification and location: {str(e)}", exc_info=True)

        return verification_results, phrase_locations

    def add_annotations(
        self, phrase_locations: Dict[str, List[Dict[str, Any]]]
    ) -> bytes:
        """Adds highlights to the PDF based on found phrase locations (from aggregated results)."""
        if not phrase_locations:
            logger.warning("No phrase locations provided for annotation. Returning original PDF bytes.")
            return self.pdf_bytes

        doc = None
        try:
            doc = fitz.open(stream=self.pdf_bytes, filetype="pdf")
            annotated_count = 0
            highlight_color = [1, 0.9, 0.3]  # Yellow
            fallback_color = [0.5, 0.7, 1.0]  # Light Blue for fallback

            # Flatten all locations from the dict for easier processing
            all_locs = []
            for phrase, locations in phrase_locations.items():
                for loc in locations:
                    # Add the phrase back into the location dict for context in annotation info
                    loc['phrase_text'] = phrase
                    all_locs.append(loc)

            # Optional: Sort annotations to potentially process page by page
            # all_locs.sort(key=lambda x: (x.get('page_num', -1), x.get('rect', [0,0,0,0])[1]))

            for loc in all_locs:
                try:
                    page_num = loc.get("page_num")
                    rect_coords = loc.get("rect")
                    method = loc.get("method", "unknown")
                    phrase = loc.get("phrase_text", "Unknown Phrase")

                    if page_num is None or rect_coords is None:
                        logger.warning(f"Skipping annotation due to missing page_num/rect for phrase '{phrase[:50]}...': {loc}")
                        continue

                    if 0 <= page_num < doc.page_count:
                        page = doc[page_num]
                        rect = fitz.Rect(rect_coords)
                        if not rect.is_empty:
                            color = fallback_color if "fallback" in method else highlight_color
                            highlight = page.add_highlight_annot(rect)
                            highlight.set_colors(stroke=color)
                            highlight.set_info(
                                content=(f"Verified ({method}, Score: {loc.get('match_score', 'N/A'):.0f}): {phrase[:100]}...")
                            )
                            highlight.update(opacity=0.4)
                            annotated_count += 1
                        # else: logger.debug(f"Skipping annotation for empty rect: {rect}")
                    # else: logger.warning(f"Skipping annotation due to invalid page num {page_num} from location data.")
                except Exception as annot_err:
                    logger.error(f"Error adding annotation for phrase '{phrase[:50]}...' at {loc}: {annot_err}")

            if annotated_count > 0:
                logger.info(f"Added {annotated_count} highlight annotations.")
                annotated_bytes = doc.tobytes(garbage=4, deflate=True)
            else:
                logger.warning("No annotations were successfully added. Returning original PDF bytes.")
                annotated_bytes = self.pdf_bytes

            return annotated_bytes

        except Exception as e:
            logger.error(f"Failed to add annotations: {str(e)}", exc_info=True)
            return self.pdf_bytes # Return original on error
        finally:
            if doc: doc.close()

class WordProcessor:
    """Handles Word document conversion to PDF."""

    def __init__(self, docx_bytes: bytes):
        if not isinstance(docx_bytes, bytes):
            raise ValueError("docx_bytes must be of type bytes")
        self.docx_bytes = docx_bytes
        logger.info(f"WordProcessor initialized with {len(docx_bytes)} bytes.")

    def convert_to_pdf_bytes(self) -> Optional[bytes]:
        """Converts the DOCX to PDF using a basic text-dumping approach."""
        logger.warning("Using basic DOCX to PDF conversion (text dump). Formatting will be lost.")
        try:
            doc = DocxDocument(BytesIO(self.docx_bytes))
            full_text = []
            for para in doc.paragraphs: full_text.append(para.text)
            # Basic table extraction (join cells with tabs, rows with newlines)
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append("\t".join(row_text))
                full_text.append("\n".join(table_text))

            extracted_text = "\n".join(full_text)

            if not extracted_text.strip():
                logger.warning("No text extracted from DOCX file.")
                return self.create_minimal_empty_pdf()

            pdf_doc = fitz.open()
            page = pdf_doc.new_page()
            rect = fitz.Rect(50, 50, page.rect.width - 50, page.rect.height - 50)
            fontsize = 10
            # Insert text with basic formatting interpretation
            res = page.insert_textbox(
                rect,
                extracted_text,
                fontsize=fontsize,
                fontname="helv", # Use a standard font
                align=fitz.TEXT_ALIGN_LEFT
            )
            if res < 0: logger.warning(f"Text might have been truncated during basic PDF creation (return code: {res}).")

            output_buffer = BytesIO()
            pdf_doc.save(output_buffer, garbage=4, deflate=True)
            pdf_doc.close()
            logger.info("Successfully created basic PDF from DOCX text.")
            return output_buffer.getvalue()

        except ImportError:
            logger.error("python-docx not installed. Cannot process Word files.")
            st.error("python-docx is required to process Word documents. Please install it (`pip install python-docx`) and restart.")
            raise Exception("python-docx is required to process Word documents.")
        except Exception as e:
            logger.error(f"Error during basic DOCX to PDF conversion: {e}", exc_info=True)
            return self.create_minimal_empty_pdf()

    def create_minimal_empty_pdf(self) -> bytes:
        """Creates a minimal valid empty PDF."""
        try:
            min_pdf = fitz.open()
            min_pdf.new_page()
            buffer = BytesIO()
            min_pdf.save(buffer)
            min_pdf.close()
            logger.warning("Created minimal empty PDF as fallback.")
            return buffer.getvalue()
        except Exception as min_e:
            logger.error(f"Failed to create minimal PDF: {min_e}")
            # Absolute last resort (often unnecessary if fitz works)
            return (b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
                    b"2 0 obj<</Type/Pages/Count 0>>endobj\nxref\n0 3\n"
                    b"0000000000 65535 f \n0000000010 00000 n \n0000000059 00000 n \n"
                    b"trailer<</Size 3/Root 1 0 R>>\nstartxref\n109\n%%EOF\n")


# --- Streamlit UI Functions ---

# --- PDF Viewer Logic ---
def update_pdf_view(pdf_bytes, page_num, filename):
    """Updates session state for the PDF viewer and triggers a rerun."""
    # Removed default values, expects args from on_click
    # if pdf_bytes is None: pdf_bytes = st.session_state.get('pdf_bytes')
    # if filename is None: filename = st.session_state.get('current_pdf_name', 'document.pdf')
    # if not isinstance(page_num, int) or page_num < 1: page_num = st.session_state.get('pdf_page', 1)

    state_changed = False
    if st.session_state.get('pdf_page') != page_num:
        st.session_state.pdf_page = page_num
        state_changed = True
    # Careful bytes comparison: only update if actually different and not None
    if pdf_bytes is not None and pdf_bytes != st.session_state.get('pdf_bytes'):
        st.session_state.pdf_bytes = pdf_bytes
        state_changed = True
    if st.session_state.get('current_pdf_name') != filename:
        st.session_state.current_pdf_name = filename
        state_changed = True
    # Always ensure PDF is shown when clicking a citation
    if not st.session_state.get('show_pdf', False):
        st.session_state.show_pdf = True
        state_changed = True

    if state_changed:
        logger.info(f"Updating PDF view state via chat citation: page={page_num}, filename={filename}, show={st.session_state.show_pdf}")
        st.rerun() # Trigger rerun to update the PDF viewer UI
    else:
        # If state didn't change but PDF wasn't visible, ensure rerun
        if not st.session_state.get('show_pdf', False):
            st.session_state.show_pdf = True
            logger.info(f"Showing PDF viewer for citation click (no state change): page={page_num}, filename={filename}")
            st.rerun()

def display_pdf_viewer():
    """Renders the PDF viewer based on session state."""
    pdf_bytes = st.session_state.get("pdf_bytes")
    show_pdf = st.session_state.get("show_pdf", False)
    current_page = st.session_state.get("pdf_page", 1)
    filename = st.session_state.get("current_pdf_name", "PDF Viewer")

    # Use expander's default state based on show_pdf
    with st.expander("📄 PDF Viewer", expanded=show_pdf):
        if not pdf_bytes:
            st.info("Upload and process a file, or click a citation to view the PDF here.")
            return

        fitz_doc = None
        try:
            fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            total_pages = fitz_doc.page_count
            if total_pages == 0:
                st.warning("The PDF document appears to have 0 pages.")
                return

            # --- Navigation ---
            st.caption(f"**{filename}**")
            cols = st.columns([1, 3, 1])

            with cols[0]:
                prev_key = f"pdf_prev_{filename}_{total_pages}" # More unique key
                if st.button("⬅️", key=prev_key, help="Previous Page", disabled=(current_page <= 1)):
                    st.session_state.pdf_page = max(1, current_page - 1)
                    st.rerun()

            with cols[1]:
                nav_key = f"nav_{filename}_{total_pages}"
                selected_page = st.number_input(
                    "Page", min_value=1, max_value=total_pages, value=current_page, step=1,
                    key=nav_key, label_visibility="collapsed", help=f"Enter page number (1-{total_pages})"
                )
                if selected_page != current_page:
                    if 1 <= selected_page <= total_pages:
                        st.session_state.pdf_page = selected_page
                        st.rerun()

            with cols[2]:
                next_key = f"pdf_next_{filename}_{total_pages}" # More unique key
                if st.button("➡️", key=next_key, help="Next Page", disabled=(current_page >= total_pages)):
                    st.session_state.pdf_page = min(total_pages, current_page + 1)
                    st.rerun()

            st.caption(f"Page {current_page} of {total_pages}")

            # --- Render Page using Fitz ---
            try:
                page = fitz_doc.load_page(current_page - 1) # 0-indexed
                pix = page.get_pixmap(dpi=150) # Adjust DPI
                img_bytes = pix.tobytes("png")
                st.image(img_bytes, use_container_width=True)
            except Exception as fitz_render_err:
                 logger.error(f"Error rendering page {current_page} with Fitz: {fitz_render_err}")
                 st.error(f"Error displaying page {current_page}.")

        except fitz.fitz.FileNotFoundError: # More specific type
            st.error("Failed to open the PDF data. It might be corrupted or empty.")
            st.session_state.show_pdf = False # Hide viewer on error
        except Exception as e:
            logger.error(f"Error displaying PDF viewer: {e}", exc_info=True)
            st.error(f"An error occurred while displaying the PDF: {e}")
            st.session_state.show_pdf = False # Hide viewer on critical error
        finally:
            if fitz_doc: fitz_doc.close()


# --- Analysis Display Logic ---
def find_best_location(locations: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """Finds the best location (prioritizing exact matches, then highest score, then earliest page)."""
    if not locations: return None

    def sort_key(loc):
        method = loc.get("method", "unknown")
        score = loc.get("match_score", 0)
        page_num = loc.get("page_num", 9999)
        method_priority = {"exact_cleaned_search": 0, "fuzzy_chunk_fallback": 1}.get(method, 99)
        return (method_priority, -score, page_num)

    valid_locations = [loc for loc in locations if "page_num" in loc and "rect" in loc]
    if not valid_locations: return None
    valid_locations.sort(key=sort_key)
    return valid_locations[0]

def display_analysis_results(analysis_results: List[Dict[str, Any]]):
    """Displays the aggregated analysis sections and citations using tabs for each file."""
    if not analysis_results:
        logger.info("No analysis results to display.")
        return

    # Add custom styling for the analysis display
    st.markdown("""
    <style>
    /* Common styles for containers */
    .sleek-container {
        background-color: #f5f5f5;
        border-radius: 8px;
        padding: 8px 16px;
        margin: 0 0 16px 0;
        display: flex;
        align-items: center;
        justify-content: space-between;
        border: 1px solid #e0e0e0;
    }
    
    /* Header title styling */
    .header-title {
        font-weight: 700;
        font-size: 1.5rem;
        color: #333;
        margin: 0;
        padding: 0;
    }
    
    /* File info container */
    .file-info-container {
        display: flex;
        align-items: center;
        padding: 0;
        margin: 0;
    }
    
    /* File name styling */
    .file-name {
        font-weight: 600;
        color: #424242;
        font-size: 1rem;
        display: flex;
        align-items: center;
        margin: 0;
        padding: 0;
    }
    
    /* File icon styling */
    .file-icon {
        color: #1976d2;
        margin-right: 8px;
    }
    
    /* Button container */
    .button-container {
        display: flex;
        gap: 8px;
    }
    
    /* Adjust button vertical alignment and size */
    .stButton > button {
        margin-top: 0 !important;
        padding-top: 2px !important;
        padding-bottom: 2px !important;
        line-height: 1.2 !important;
    }
    
    /* Remove extra padding from Streamlit containers */
    .block-container {
        padding-top: 0 !important;
        padding-bottom: 0 !important;
    }
    
    /* Reduce padding in st.container */
    .css-ocqkz7 {
        padding: 0 !important;
    }
    
    /* Custom container for buttons */
    .button-row {
        display: flex;
        flex-direction: row;
        justify-content: flex-end;
        align-items: center;
        gap: 10px;
        width: 100%;
        flex-wrap: nowrap;
    }
    
    /* Make buttons more compact */
    .stButton > button {
        margin-top: 0 !important;
        margin-bottom: 0 !important;
        padding: 2px 8px !important;
        line-height: 1.2 !important;
        white-space: nowrap !important;
        font-size: 0.9rem !important;
    }
    </style>
    """, unsafe_allow_html=True)

    analysis_col, pdf_col = st.columns([2.5, 1.5], gap="small")

    with analysis_col:
        # Create two columns for the header layout
        header_col1, header_col2 = st.columns([0.6, 0.4])
        
        # Add the title in the first column
        with header_col1:
            st.markdown('<div class="header-title">AI Analysis Results</div>', unsafe_allow_html=True)
        
        # Add the title in the second column (empty for now)
        with header_col2:
            # Empty space for alignment
            st.write("")
        
        # Add a thin divider after the header section
        st.markdown('<hr style="margin: 12px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

        # Filter results first
        success_results = [
            r for r in analysis_results
            if isinstance(r, dict) and "error" not in r and "ai_analysis" in r
        ]
        error_results = [
            r for r in analysis_results
            if isinstance(r, dict) and "error" in r
        ]

        # Check if any successful results contain actual analysis data
        results_with_real_analysis = []
        for r in success_results:
             try:
                 parsed_ai_analysis = json.loads(r["ai_analysis"])
                 if isinstance(parsed_ai_analysis, dict) and "analysis_sections" in parsed_ai_analysis:
                     sections = parsed_ai_analysis["analysis_sections"]
                     if isinstance(sections, dict) and any(not k.startswith(("error_", "info_", "skipped_")) for k in sections):
                         results_with_real_analysis.append((r, parsed_ai_analysis)) # Store tuple of result and parsed data
                         continue # Move to next result
             except json.JSONDecodeError:
                 logger.warning(f"Failed to parse AI analysis JSON for result display: {r.get('filename', 'Unknown')}")
                 # Treat as error for display purposes if JSON is invalid
                 error_results.append(r)
             except Exception as e:
                  logger.error(f"Unexpected error checking analysis data for {r.get('filename', 'Unknown')}: {e}", exc_info=True)
                  error_results.append(r)

        if not results_with_real_analysis:
            st.info(
                "Processing complete, but no analysis sections were successfully generated or found "
                "in the successful results. Check logs or errors below."
            )
            # Display errors/info from original error_results
            if error_results:
                with st.expander("Processing Errors/Info Summary", expanded=True):
                    for err_res in error_results:
                         st.warning(f"**{err_res.get('filename', 'File')}**: {err_res.get('error', 'No details') or err_res.get('ai_analysis', 'No details')}")
        else:
            # --- Create Tabs for Each Successful Analysis --- 
            tab_titles = [res[0].get("filename", f"Result {i+1}") for i, res in enumerate(results_with_real_analysis)]
            tabs = st.tabs(tab_titles)

            for i, (result, ai_analysis) in enumerate(results_with_real_analysis):
                with tabs[i]:
                    filename = result.get("filename", f"Result {i+1}")
                    verification_results = result.get("verification_results", {})
                    phrase_locations = result.get("phrase_locations", {})
                    annotated_pdf_b64 = result.get("annotated_pdf")
                    
                    # Create a sleeker container for file information
                    # Create two columns for the file info layout
                    file_col1, file_col2 = st.columns([0.8, 0.2])
                    
                    # Add the filename in the first column
                    with file_col1:
                        st.markdown(f"""
                        <div class="sleek-container" style="margin-bottom: 16px;">
                            <div class="file-name">
                                <span class="file-icon">📄</span> {filename}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Add download button in the second column
                    with file_col2:
                        if annotated_pdf_b64:
                            try:
                                annotated_pdf_bytes = base64.b64decode(annotated_pdf_b64)
                                # Determine if this is a PDF or Word document
                                is_word_doc = filename.lower().endswith('.docx')
                                
                                # Set appropriate label and file extension
                                if is_word_doc:
                                    download_label = "Download Annotated PDF"
                                    file_extension = '.pdf'
                                else:
                                    download_label = "Download Annotated PDF"
                                    file_extension = '.pdf'
                                
                                # Add download button
                                st.download_button(
                                    label="💾 PDF",
                                    data=annotated_pdf_bytes,
                                    file_name=f"{filename.replace('.pdf', '').replace('.docx', '')}_annotated{file_extension}",
                                    mime="application/pdf",
                                    key=f"download_{filename}_{i}",
                                    help=f"Download annotated PDF for {filename}",
                                    use_container_width=True,
                                )
                            except Exception as decode_err:
                                logger.error(f"Failed to decode annotated PDF for download button ({filename}): {decode_err}")
                                st.caption("DL Err")
                        else:
                            st.caption("No PDF")

                    try:
                        # Use the already parsed ai_analysis dictionary 
                        analysis_sections = ai_analysis.get("analysis_sections", {})
                        if not analysis_sections:
                            st.warning("No analysis sections found in the AI response for this file.")
                            continue # Skip to next tab

                        citation_counter = 0
                        for section_key, section_data in analysis_sections.items():
                            if not isinstance(section_data, dict): continue

                            if section_key.startswith(("error_", "info_", "skipped_")):
                                st.caption(f"Skipped/Error Section: '{section_key}' - Check logs or error summary.")
                                continue

                            display_section_name = section_key.replace("_", " ").title()

                            # Create a container for the section title with improved styling
                            with st.container(border=False):
                                st.markdown(f"""
                                    <div style='background-color: #f5f5f5; padding: 0px 16px; border-radius: 8px; 
                                            margin: 16px 0 8px 0; border-left: 4px solid #1976d2;'>
                                        <h4 style='color: #333; font-size: 1.2rem; margin: 0; font-weight: 600;'>
                                            {display_section_name}
                                        </h4>
                                    </div>
                                """, unsafe_allow_html=True)
                            
                            # Section content in a bordered container
                            with st.container(border=True):
                                if section_data.get("Analysis"):
                                    # Create a container for both analysis and context
                                    with st.container():
                                        # Render the analysis section with context included directly
                                        analysis_html = f"""
                                            <div style='background-color: #f8f9fa; padding: .5rem; border-radius: 0.5rem; margin-bottom: 1rem;'>
                                                <h4 style='color: #1e88e5; font-size: 1.1rem;'>
                                                    Analysis
                                                </h4>
                                                <div style='color: #424242; line-height: 1.6;'>
                                                    {section_data["Analysis"]}"""
                                        
                                        # Add context directly in the analysis div if it exists
                                        if section_data.get("Context"):
                                            analysis_html += f"""
                                                    <div style='margin-top: 0.8rem; border-top: 1px solid #e0e0e0; padding-top: 0.8rem;'>
                                                        <span style='color: #1b5e20; font-size: 0.9rem; line-height: 1.4;'>{section_data.get("Context", "")}</span>
                                                    </div>"""
                                        
                                        # Close the divs and render the HTML
                                        analysis_html += """
                                                </div>
                                            </div>
                                        """
                                        st.markdown(analysis_html, unsafe_allow_html=True)

                                supporting_phrases = section_data.get("Supporting_Phrases", section_data.get("supporting_quotes", []))
                                if not isinstance(supporting_phrases, list): supporting_phrases = []

                                if supporting_phrases:
                                    # --- Determine if the expander should be open by default ---
                                    expand_citations = False
                                    for phrase_text in supporting_phrases:
                                        if isinstance(phrase_text, str) and phrase_text != "No relevant phrase found.":
                                            is_verified = verification_results.get(phrase_text, False)
                                            if not is_verified:
                                                expand_citations = True
                                                break # Found one unverified, no need to check further

                                    # --- Create the expander with conditional expansion ---
                                    with st.expander("Supporting Citations", expanded=expand_citations):
                                        has_citations_to_show = False
                                        for phrase_idx, phrase_text in enumerate(supporting_phrases):
                                            if not isinstance(phrase_text, str) or phrase_text == "No relevant phrase found.":
                                                continue
                                            has_citations_to_show = True
                                            citation_counter += 1
                                            is_verified = verification_results.get(phrase_text, False)
                                            locations = phrase_locations.get(phrase_text, [])
                                            best_location = find_best_location(locations)

                                            score_info = f"Score: {best_location['match_score']:.1f}" if best_location and "match_score" in best_location else ""
                                            method_info = f"{best_location['method']}" if best_location and "method" in best_location else ""
                                            page_info = f"Pg {best_location['page_num'] + 1}" if best_location and "page_num" in best_location else ""

                                            cite_col, btn_col = st.columns([0.90, 0.10], gap="small")
                                            with cite_col:
                                                if is_verified:
                                                    badge_html = '<span style="display: inline-block; background-color: #d1fecf; color: #11631a; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #a1e0a3; font-weight: 600;">✔ Verified</span>'
                                                else:
                                                    badge_html = '<span style="display: inline-block; background-color: #ffeacc; color: #a05e03; padding: 1px 6px; border-radius: 0.25rem; font-size: 0.8em; margin-left: 5px; border: 1px solid #f8c78d; font-weight: 600;">⚠️ Needs Review</span>'

                                                st.markdown(f"""
                                                    <div style="border: 1px solid #e0e0e0; border-radius: 5px; padding: 8px 12px; margin-top: 5px; margin-bottom: 8px; background-color: #f9f9f9;">
                                                        <div style="margin-bottom: 5px; display: flex; justify-content: space-between; align-items: center;">
                                                            <span style="font-weight: bold;">Citation {citation_counter} {badge_html}</span>
                                                            <span style="font-size: 0.8em; color: #555;">{page_info} {score_info} <span title='{method_info}'>({method_info})</span></span>
                                                        </div>
                                                        <div style="color: #333; line-height: 1.4; font-size: 0.95em;"><i>"{phrase_text}"</i></div>
                                                    </div>""", unsafe_allow_html=True)
                                            with btn_col:
                                                if is_verified and best_location and "page_num" in best_location and annotated_pdf_b64:
                                                    page_num_1_indexed = best_location["page_num"] + 1
                                                    button_key = f"goto_{filename}_{section_key}_{citation_counter}_{i}_{phrase_idx}" # More unique key
                                                    st.markdown('<div style="margin-top: 20px;"></div>', unsafe_allow_html=True)
                                                    if st.button("Go", key=button_key, type="secondary", help=f"Go to Page {page_num_1_indexed} in {filename}"):
                                                        try:
                                                            pdf_bytes = base64.b64decode(annotated_pdf_b64)
                                                            update_pdf_view(pdf_bytes=pdf_bytes, page_num=page_num_1_indexed, filename=filename)
                                                            # No rerun needed here, update_pdf_view handles it if necessary
                                                        except Exception as decode_err:
                                                            logger.error(f"Failed to decode/set PDF for citation button: {decode_err}", exc_info=True)
                                                            st.warning("Could not load PDF for this citation.")
                                                elif is_verified:
                                                    st.markdown('<div style="margin-top: 20px; text-align: center;">', unsafe_allow_html=True)
                                                    st.caption("Loc N/A")
                                                    st.markdown("</div>", unsafe_allow_html=True)
                                        if not has_citations_to_show:
                                            st.caption("No supporting citations provided or found for this section.")

                    except Exception as display_err:
                        logger.error(f"Error displaying analysis tab for {filename}: {display_err}", exc_info=True)
                        st.error(f"Error displaying analysis results for {filename}: {display_err}")

    # --- PDF Viewer and Tools Column --- 
    with pdf_col:
        # Update this to use header-title styling to match AI Analysis Results
        st.markdown('<div class="header-title">Analysis Tools & PDF Viewer</div>', unsafe_allow_html=True)
        
        # Add a thin divider after the header section
        st.markdown('<hr style="margin: 12px 0; border: 0; border-top: 1px solid #e0e0e0;">', unsafe_allow_html=True)

        # --- Wrap Tool Expanders in a Container --- 
        with st.container():
            # --- Chat Interface Expander --- 
            with st.expander("SmartChat", expanded=False):
                if not st.session_state.get("preprocessed_data"):
                    st.info("Upload and process documents to enable chat.")
                else:
                    chat_container = st.container(height=400, border=True)
                    with chat_container:
                        # Use enumerate to get the index of each message in the session state list
                        for msg_idx, message in enumerate(st.session_state.chat_messages):
                            with st.chat_message(message["role"]):
                                if message["role"] == "assistant":
                                    processed_text = message.get("processed_text", message["content"])
                                    citation_details = message.get("citation_details", [])
                                    # Pass the message index (msg_idx) to the display function
                                    display_chat_message_with_citations(processed_text, citation_details, msg_idx)
                                else:
                                    st.markdown(message["content"])

                    if prompt := st.chat_input("Ask about the uploaded documents...", key="chat_input_main"):
                        st.session_state.chat_messages.append({"role": "user", "content": prompt})
                        processed_chat_text = "Error: Could not generate response."
                        chat_citation_details = []
                        raw_ai_response_content = ""
                        try:
                            with st.spinner("Thinking..."):
                                logger.info(f"Chat RAG started for: {prompt[:50]}...")
                                CHAT_RAG_TOP_K_PER_DOC = 5 # Increased from 3 to 5, to get more context
                                relevant_chunks = retrieve_relevant_chunks_for_chat(prompt, top_k_per_doc=CHAT_RAG_TOP_K_PER_DOC)
                                analyzer = DocumentAnalyzer()
                                logger.info(f"Generating chat response for: {prompt[:50]}...")
                                raw_ai_response_content = run_async(analyzer.generate_chat_response(prompt, relevant_chunks))
                                logger.info("Chat response generated.")
                                processed_chat_text, chat_citation_details = process_chat_response_for_numbered_citations(raw_ai_response_content)
                        except Exception as chat_err:
                            logger.error(f"Error during chat processing: {chat_err}", exc_info=True)
                            processed_chat_text = f"Sorry, an error occurred while processing your request: {str(chat_err)}"
                            chat_citation_details = []
                        st.session_state.chat_messages.append({
                            "role": "assistant",
                            "content": raw_ai_response_content,
                            "processed_text": processed_chat_text,
                            "citation_details": chat_citation_details
                        })
                        st.rerun()

            # --- Export Expander --- 
            with st.expander("Export Results", expanded=False):
                # Use results_with_real_analysis which contains successful results with content
                if results_with_real_analysis:
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    try:
                        flat_data = []
                        exportable_results_list = [res[0] for res in results_with_real_analysis] # Extract original result dicts
                        parsed_analysis_list = [res[1] for res in results_with_real_analysis] # Extract parsed analysis dicts

                        for idx, exportable_result in enumerate(exportable_results_list):
                            fname = exportable_result.get("filename", "N/A")
                            ai_data = parsed_analysis_list[idx] # Get corresponding parsed data
                            title = ai_data.get("title", "")
                            verif_res = exportable_result.get("verification_results", {})
                            phrase_locs = exportable_result.get("phrase_locations", {})

                            for sec_name, sec_data in ai_data.get("analysis_sections", {}).items():
                                if not isinstance(sec_data, dict) or sec_name.startswith(("error_", "info_", "skipped_")):
                                    continue
                                analysis = sec_data.get("Analysis", "")
                                context = sec_data.get("Context", "")
                                phrases = sec_data.get("Supporting_Phrases", sec_data.get("supporting_quotes", []))
                                if not isinstance(phrases, list): phrases = []

                                if not phrases or phrases == ["No relevant phrase found."]:
                                    flat_data.append({"Filename": fname, "AI Title": title, "Section": sec_name, "Analysis": analysis, "Context": context, "Supporting Phrase": "N/A", "Verified": "N/A", "Page": "N/A", "Match Score": "N/A", "Method": "N/A"})
                                else:
                                    for phrase in phrases:
                                        if not isinstance(phrase, str): continue
                                        verified = verif_res.get(phrase, False)
                                        locs = phrase_locs.get(phrase, [])
                                        best_loc = find_best_location(locs)
                                        page = best_loc["page_num"] + 1 if best_loc and "page_num" in best_loc else "N/A"
                                        score = f"{best_loc['match_score']:.1f}" if best_loc and "match_score" in best_loc else "N/A"
                                        method = best_loc["method"] if best_loc and "method" in best_loc else "N/A"
                                        flat_data.append({"Filename": fname, "AI Title": title, "Section": sec_name, "Analysis": analysis, "Context": context, "Supporting Phrase": phrase, "Verified": verified, "Page": page, "Match Score": score, "Method": method})

                        if not flat_data:
                            st.info("No data available to export after filtering analysis sections.")
                        else:
                            col1, col2 = st.columns(2)
                            df = pd.DataFrame(flat_data)

                            # Export to Excel 
                            excel_buffer = BytesIO()
                            try:
                                df.to_excel(excel_buffer, index=False, engine="openpyxl")
                                excel_buffer.seek(0)
                                with col1: st.download_button(
                                    "📥 Export Excel",
                                    excel_buffer,
                                    f"analysis_report_{timestamp}.xlsx", # Generic name for multi-file export
                                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key="export_excel_main"
                                )
                            except ImportError:
                                 logger.error("Export to Excel failed: 'openpyxl' not found.")
                                 with col1: st.warning("Excel export requires 'openpyxl'. Install it (`pip install openpyxl`)")
                            except Exception as excel_export_err:
                                logger.error(f"Excel export failed: {excel_export_err}", exc_info=True)
                                with col1: st.error(f"Excel export error: {excel_export_err}")

                            # Export to Word 
                            try:
                                # Pass the list of successful, analyzed results
                                word_bytes = export_to_word(exportable_results_list)
                                with col2: st.download_button(
                                    "📥 Export Word",
                                    word_bytes,
                                    f"analysis_report_{timestamp}.docx", # Generic name
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="export_word_main"
                                )
                            except ImportError:
                                 logger.error("Export to Word failed: 'python-docx' not found.")
                                 with col2: st.warning("Word export requires 'python-docx'. Install it (`pip install python-docx`)")
                            except Exception as word_export_err:
                                logger.error(f"Word export failed: {word_export_err}", exc_info=True)
                                with col2: st.error(f"Word export error: {word_export_err}")

                    except Exception as export_setup_err:
                        logger.error(f"Export data preparation failed: {export_setup_err}", exc_info=True)
                        st.error(f"Failed to prepare data for export: {export_setup_err}")
                else:
                    st.info("No analysis results available to export.")

            # --- Report Issue Expander --- 
            with st.expander("⚠️ Report Issue", expanded=False):
                st.markdown("""
                    ### Report an Issue
                    If you encounter an issue, please describe it below, download the report package, and then **manually email the package** to CNT_Automations@ifc.org.
                """)
                
                # Issue description
                issue_description = st.text_area(
                    "Describe the issue",
                    placeholder="Please describe what went wrong or what results were inaccurate...",
                    height=100,
                    key="issue_description_input" # Added a key
                )

                report_package_bytes = None
                report_filename = f'issue_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip'

                # --- Inner function to create the report package --- 
                # ... (create_report_package_content function remains the same) ...
                def create_report_package_content(desc):
                    try:
                        report_data = {
                            "timestamp": datetime.now().isoformat(),
                            "issue_description": desc, # Use passed description
                            "user_inputs": {
                                "prompt": st.session_state.get('user_prompt', ''),
                                # Add safe .get() for potentially missing keys
                                "threshold": st.session_state.get('threshold', None),
                                "keywords_input": st.session_state.get('keywords_input', None),
                                "generated_keywords": st.session_state.get('generated_keywords', None)
                            },
                            # Make sure analysis_results is serializable (it should be if it's based on JSON)
                            "analysis_results": st.session_state.get('analysis_results', None),
                            "current_document": st.session_state.get('current_pdf_name', None),
                            "preprocessed_data_keys": list(st.session_state.get('preprocessed_data', {}).keys()),
                            "chat_history_summary": [
                                {"role": msg.get("role"), "content_preview": msg.get("content", "")[:100]+"..."}
                                for msg in st.session_state.get("chat_messages", [])
                            ]
                        }
                        
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            # Write report data as JSON
                            try:
                                zip_file.writestr('report_data.json', json.dumps(report_data, indent=2, default=str)) # Add default=str for non-serializable types
                            except Exception as json_err:
                                zip_file.writestr('report_data_error.txt', f"Error serializing report data: {json_err}")
                                logger.error(f"Error serializing report_data.json: {json_err}", exc_info=True)

                            # Write original uploaded files
                            uploaded_file_objs = st.session_state.get('uploaded_file_objects') # Use correct key
                            if uploaded_file_objs:
                                for uploaded_file in uploaded_file_objs:
                                    try:
                                        # Ensure the file object is valid and readable
                                        if hasattr(uploaded_file, 'name') and hasattr(uploaded_file, 'getvalue'):
                                            zip_file.writestr(f'original_docs/{uploaded_file.name}', uploaded_file.getvalue())
                                        else:
                                             logger.warning(f"Skipping invalid file object in uploaded_file_objects during report creation: {type(uploaded_file)}")
                                    except Exception as file_read_err:
                                        zip_file.writestr(f'original_docs/ERROR_{uploaded_file.name}.txt', f"Error reading file: {file_read_err}")
                                        logger.error(f"Error reading file {uploaded_file.name} for report package: {file_read_err}", exc_info=True)

                            # Write annotated PDFs
                            analysis_results_list = st.session_state.get('analysis_results')
                            if analysis_results_list:
                                for result in analysis_results_list:
                                    if isinstance(result, dict) and 'annotated_pdf' in result and result.get('annotated_pdf'):
                                        try:
                                            pdf_bytes = base64.b64decode(result['annotated_pdf'])
                                            pdf_filename = result.get('filename', f'unknown_annotated_{result.get("timestamp", "ts")}.pdf')
                                            zip_file.writestr(f'annotated_pdfs/{pdf_filename}', pdf_bytes)
                                        except Exception as pdf_err:
                                             zip_file.writestr(f'annotated_pdfs/ERROR_{result.get("filename", "unknown")}.txt', f"Error decoding/writing PDF: {pdf_err}")
                                             logger.error(f"Error writing annotated PDF {result.get('filename')} to report: {pdf_err}", exc_info=True)
                        
                        zip_buffer.seek(0)
                        return zip_buffer.getvalue()
                    except Exception as zip_e:
                        logger.error(f"Error creating report package zip file: {zip_e}", exc_info=True)
                        # Create a simple error zip as fallback
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            zip_file.writestr('error_creating_report.txt', f"Failed to create full report package: {zip_e}")
                        zip_buffer.seek(0)
                        return zip_buffer.getvalue()
                # --- End of inner function ---

                # Button to download the package (now takes full width)
                download_disabled = not issue_description.strip()
                try:
                    # Generate package content only if description is provided
                    if not download_disabled:
                        report_package_bytes = create_report_package_content(issue_description)
                    else:
                        report_package_bytes = b"" # Ensure it's defined as empty bytes if disabled

                    st.download_button(
                        label="📥 Download Report Package",
                        on_click="ignore",
                        data=report_package_bytes,
                        file_name=report_filename,
                        mime='application/zip',
                        disabled=download_disabled,
                        help="Provide an issue description first, then download the package to attach to your email.",
                        key="download_report_button",
                        use_container_width=True # Make button full width
                    )
                except Exception as e:
                    st.error(f"Error preparing report package: {str(e)}")
                    logger.error(f"Error preparing report package for download button: {str(e)}", exc_info=True)

                # Removed the Mail client button/link section

        # --- PDF Viewer Display (Remains outside the container, at the END) --- 
        display_pdf_viewer()


# --- NEW: Chat UI Helper Functions ---

def find_annotated_pdf_for_filename(filename: str) -> Optional[bytes]:
    """Finds the base64 decoded annotated PDF bytes for a given filename from session state."""
    for result in st.session_state.get("analysis_results", []):
        if isinstance(result, dict) and result.get("filename") == filename and result.get("annotated_pdf"):
            try:
                return base64.b64decode(result["annotated_pdf"])
            except Exception as e:
                logger.error(f"Failed to decode annotated PDF for {filename} in chat citation: {e}")
                return None
    logger.warning(f"Could not find annotated PDF data for {filename} in session state analysis_results.")
    return None

# --- NEW: Process AI response for numbered citations ---
def process_chat_response_for_numbered_citations(raw_response_text: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Processes raw AI response text containing (Source:...) citations.
    Replaces them with sequential numbers [1], [2], etc., and returns the
    modified text along with a list of citation details for creating buttons.

    Args:
        raw_response_text: The original text from the AI.

    Returns:
        Tuple containing:
        - str: The response text with inline citations replaced by numbers ([1], [2]).
        - List[Dict[str, Any]]: A list of citation details, each dict containing 
                                 'number', 'filename', 'page', 'pdf_bytes'.
    """
    citation_pattern = re.compile(r"\(Source:\s*(?P<filename>[^,]+?)\s*,\s*Page:\s*(?P<page>\d+)\)")
    
    citations_found_for_replacement = [] # Stores info needed for text replacement
    citation_details_for_footer = [] # Stores unique details for footer buttons
    next_citation_number = 1
    processed_text = raw_response_text

    # Find all citations and assign sequential numbers
    for match in citation_pattern.finditer(raw_response_text):
        filename = match.group("filename").strip()
        page_str = match.group("page").strip()
        try:
            page_num = int(page_str)

            # --- Assign unique number to THIS instance ---
            current_number = next_citation_number

            # Get PDF bytes for this source
            pdf_bytes = find_annotated_pdf_for_filename(filename)

            # Store details for the footer button list
            citation_details_for_footer.append({
                'number': current_number,
                'filename': filename,
                'page': page_num,
                'pdf_bytes': pdf_bytes # Can be None if not found
            })

            # Store details needed to replace the text later
            citations_found_for_replacement.append({
                'start': match.start(),
                'end': match.end(),
                'number': current_number,
                'original_text': match.group(0)
            })

            # Increment for the *next* citation found
            next_citation_number += 1

        except ValueError:
            logger.warning(f"Found invalid page number in citation: {match.group(0)}")
        except Exception as e:
            logger.error(f"Error processing citation {match.group(0)}: {e}")

    # Second pass: Replace citations in the text from end to start (to avoid index issues)
    # Sort by start position in reverse order
    citations_found_for_replacement.sort(key=lambda x: x['start'], reverse=True)

    for citation in citations_found_for_replacement:
        processed_text = (
            processed_text[:citation['start']] +
            f" [{citation['number']}]" +
            processed_text[citation['end']:]
        )

    # Footer details are already collected sequentially, no extra sort needed by number
    # citation_details_for_footer.sort(key=lambda x: x['number']) # This is no longer needed

    return processed_text.strip(), citation_details_for_footer


def display_chat_message_with_citations(processed_text: str, citation_details: List[Dict[str, Any]], msg_idx: int):
    """
    Displays the processed chat message containing numbered citations [1], [2], etc.,
    and lists the corresponding source buttons below.

    Args:
        processed_text: The message text with (Source:...) replaced by [1], [2].
        citation_details: A list of dictionaries from process_chat_response_for_numbered_citations,
                          each containing 'number', 'filename', 'page', 'pdf_bytes'.
        msg_idx: The index of the message in the overall chat history (for unique keys).
    """

    # Display the main message content with inline numbers
    st.markdown(processed_text, unsafe_allow_html=True)

    # Display the citation sources below if any exist
    if citation_details:
        st.markdown("---") # Add a separator
        st.caption("Sources:")
        # Use columns for a slightly more compact layout if many sources
        num_columns = min(len(citation_details), 3) # Max 3 columns
        cols = st.columns(num_columns)
        col_idx = 0

        for i, citation in enumerate(citation_details):
            number = citation['number']
            filename = citation['filename']
            page_num = citation['page']
            pdf_bytes = citation['pdf_bytes']

            with cols[col_idx]:
                if pdf_bytes:
                    # Include the message index (msg_idx) and citation index (i) for guaranteed uniqueness
                    button_key = f"chat_footer_cite_{filename}_{page_num}_{number}_{msg_idx}_{i}" # Updated unique key
                    # Render the button with the citation number
                    st.button(f"[{number}] 📄 {filename}, p{page_num}",
                              key=button_key,
                              help=f"View Page {page_num} in {filename}",
                              type="secondary",
                              on_click=update_pdf_view,
                              args=(pdf_bytes, page_num, filename)
                              )
                else:
                    # If PDF not found, display text indicating the source
                    st.markdown(f"[{number}] {filename}, p{page_num} (PDF not found)")

            col_idx = (col_idx + 1) % num_columns # Cycle through columns


# --- Main Application Logic ---

def aggregate_analysis_results(
    individual_results: List[Tuple[str, str, str]], # List of (title, sub_prompt, analysis_json_str) tuples
    original_filename: str,
    original_prompt: str
    ) -> str: # Returns single aggregated JSON string
    """Merges JSON results from multiple sub-prompt analyses into one."""
    final_analysis = {
        "title": f"Aggregated Analysis for {original_filename}", # Default title
        "analysis_sections": {}
    }
    found_valid_title = False # Tracks if we set a better overall title

    for i, (task_title, sub_prompt, analysis_json_str) in enumerate(individual_results):
        try:
            parsed_result = json.loads(analysis_json_str)
            if not isinstance(parsed_result, dict):
                 logger.warning(f"Sub-analysis result {i} is not a dict, skipping.")
                 final_analysis["analysis_sections"][f"error_aggregation_invalid_format_{i}"] = {
                     "Analysis": f"Result from sub-analysis {i+1} (Prompt: '{sub_prompt[:50]}...') was not in the expected dictionary format.",
                     "Supporting_Phrases": ["No relevant phrase found."],
                     "Context": "Aggregation Error"
                 }
                 continue

            # Extract fields from the simplified schema returned by analyze_document
            analysis_summary = parsed_result.get("analysis_summary", "No analysis provided.")
            supporting_quotes = parsed_result.get("supporting_quotes", ["No relevant phrase found."])
            analysis_context = parsed_result.get("analysis_context", "") # Context from analysis LLM

            # --- Use the AI-generated title to create the section key ---
            # Clean up the AI-generated title to create a valid snake_case key
            section_key = re.sub(r'[^a-zA-Z0-9]', '_', task_title)
            section_key = re.sub(r'_+', '_', section_key)  # Replace multiple underscores
            section_key = section_key.strip('_')[:60]  # Limit length and remove leading/trailing

            if not section_key: # Fallback if title was empty or only symbols
                section_key = f"sub_prompt_{i+1}"

            # Add a numeric suffix if the key already exists (handles duplicate titles)
            base_key = section_key
            counter = 1
            while section_key in final_analysis["analysis_sections"]:
                section_key = f"{base_key}_{counter}"
                counter += 1

            # Create the section with the proper format expected by the display logic
            # Include the original sub-prompt in the context for clarity
            final_analysis["analysis_sections"][section_key] = {
                "Analysis": analysis_summary,
                "Supporting_Phrases": supporting_quotes,
                "Context": f"{analysis_context}"
            }

            # Set a better overall title if we haven't yet and this is the first valid analysis
            if not found_valid_title and analysis_summary and "error" not in analysis_summary.lower() and "skipped" not in analysis_summary.lower():
                final_analysis["title"] = f"Analysis of {original_filename}"
                found_valid_title = True

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON for sub-analysis result {i} (Prompt: '{sub_prompt[:50]}...'): {e}")
            final_analysis["analysis_sections"][f"error_aggregation_json_decode_{i}"] = {
                "Analysis": f"Failed to parse JSON output from sub-analysis {i+1} (Prompt: '{sub_prompt[:50]}...'). Error: {e}",
                "Supporting_Phrases": ["No relevant phrase found."],
                "Context": "Aggregation Error"
            }
        except Exception as e:
            logger.error(f"Unexpected error aggregating sub-analysis result {i} (Prompt: '{sub_prompt[:50]}...'): {e}", exc_info=True)
            final_analysis["analysis_sections"][f"error_aggregation_unexpected_{i}"] = {
                "Analysis": f"Unexpected error processing result from sub-analysis {i+1} (Prompt: '{sub_prompt[:50]}...'). Error: {e}",
                "Supporting_Phrases": ["No relevant phrase found."],
                "Context": "Aggregation Error"
            }

    # If no valid title was found after aggregation, adjust the default
    if not found_valid_title and not final_analysis["analysis_sections"]:
        final_analysis["title"] = f"Analysis Failed for {original_filename}"
    elif not found_valid_title:
         # Keep the default or maybe indicate issues
         final_analysis["title"] = f"Aggregated Analysis for {original_filename} (Check Sections)"


    return json.dumps(final_analysis, indent=2)

# --- NEW: Word Export Function ---
def export_to_word(analysis_results):
    """Export analysis results to Word document"""
    try:
        # Create a new Word document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "Document Analysis Report"
        doc.core_properties.author = "AI Document Analyzer"
        doc.core_properties.created = datetime.now()

        # Add title
        title = doc.add_heading('Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True

        doc.add_paragraph()  # Add some space

        # Process each result
        for result in analysis_results:
            filename = result.get('filename', 'Unknown File') # Use .get for safety

            # Add document header
            doc.add_heading(f"Document: {filename}", 1)

            try:
                # Parse the analysis data with error handling
                ai_analysis_value = result.get('ai_analysis') # Use .get
                if isinstance(ai_analysis_value, str):
                    try:
                        analysis_data = json.loads(ai_analysis_value)
                    except json.JSONDecodeError as e:
                        logger.error(f"JSON parsing error in export_to_word for {filename}: {str(e)}")
                        raise ValueError(f"Could not parse analysis JSON for {filename}. Content: {ai_analysis_value[:200]}...")
                elif isinstance(ai_analysis_value, dict):
                    analysis_data = ai_analysis_value # Already a dict
                else:
                    raise ValueError(f"Unexpected type for ai_analysis in {filename}: {type(ai_analysis_value)}")

                analysis_sections = analysis_data.get('analysis_sections', {}) # Use .get

                # Check if this is an error response (handle potential errors from aggregation)
                is_error_only = True
                for section_name in analysis_sections:
                    if not section_name.startswith(("error_", "info_", "skipped_")):
                        is_error_only = False
                        break

                if not analysis_sections or is_error_only:
                    doc.add_heading("Analysis Status", 2)
                    p = doc.add_paragraph()
                    status_run = p.add_run("No analysis sections found or only errors/info were present.")
                    status_run.font.color.rgb = RGBColor(128, 128, 128) # Gray color
                    # Optionally list the specific errors/info keys found
                    if analysis_sections:
                        for error_key, error_data in analysis_sections.items():
                             if isinstance(error_data, dict):
                                  p_err = doc.add_paragraph()
                                  p_err.add_run(f"- {error_key}: ").bold = True
                                  p_err.add_run(error_data.get('Analysis', 'Details unavailable.'))
                                  if error_data.get('Context'):
                                       p_err.add_run(f" (Context: {error_data.get('Context')})")
                    doc.add_paragraph()  # Add space
                    if len(analysis_results) > 1 : doc.add_page_break() # Add page break only if more than one doc
                    continue # Skip to the next document result

                # Process each valid section in the analysis
                for section_name, section in analysis_sections.items():
                    # Skip error/info sections in the main processing loop
                    if section_name.startswith(("error_", "info_", "skipped_")):
                        continue
                    if not isinstance(section, dict): # Skip non-dict sections
                        logger.warning(f"Skipping non-dictionary section '{section_name}' during Word export for {filename}.")
                        continue

                    # Add section header
                    doc.add_heading(section_name.replace('_', ' ').title(), 2)

                    # Add analysis text
                    if section.get('Analysis'):
                        p = doc.add_paragraph()
                        p.add_run("Analysis: ").bold = True
                        p.add_run(section.get('Analysis'))

                    # Add context if available
                    if section.get('Context'):
                        p = doc.add_paragraph()
                        p.add_run("Context: ").bold = True
                        p.add_run(section.get('Context'))

                    # Add supporting phrases
                    # Use the correct key "Supporting_Phrases" from aggregation
                    supporting_phrases = section.get('Supporting_Phrases', [])
                    if supporting_phrases:
                        doc.add_heading("Supporting Citations", 3)

                        if not isinstance(supporting_phrases, list):
                             p = doc.add_paragraph()
                             p.style = 'List Bullet'
                             p.add_run(f"Warning: Expected a list for Supporting_Phrases, got {type(supporting_phrases)}.")
                             continue

                        for phrase in supporting_phrases:
                            if not isinstance(phrase, str):
                                p = doc.add_paragraph()
                                p.style = 'List Bullet'
                                p.add_run(f"Warning: Found non-string item in Supporting_Phrases: {type(phrase)}.")
                                continue

                            if phrase == "No relevant phrase found.":
                                p = doc.add_paragraph()
                                p.style = 'List Bullet'
                                p.add_run("No relevant phrase found")
                            else:
                                # Phrase text is already clean
                                clean_phrase = phrase

                                # Check verification status using the clean phrase
                                is_verified = result.get("verification_results", {}).get(clean_phrase, False)

                                # Get page number and other details from phrase_locations
                                page_num_info = "N/A"
                                score_info = "N/A"
                                method_info = "N/A"
                                locations = result.get("phrase_locations", {}).get(clean_phrase, [])
                                best_location = find_best_location(locations) # Use the helper function
                                if best_location:
                                    if "page_num" in best_location: page_num_info = str(best_location["page_num"] + 1) # 1-indexed
                                    if "match_score" in best_location: score_info = f"{best_location['match_score']:.1f}"
                                    if "method" in best_location: method_info = best_location["method"]

                                p = doc.add_paragraph()
                                p.style = 'List Bullet'

                                # Add verification icon
                                if is_verified:
                                    p.add_run("✓ ").bold = True
                                else:
                                    # Use a different symbol if not verified but location exists vs completely not found
                                    symbol = "❓ " if locations else "❌ "
                                    p.add_run(symbol).bold = True

                                # Add the phrase
                                p.add_run(clean_phrase)

                                # Add details (page, score, method)
                                details_run = p.add_run(f" (Pg: {page_num_info}, Score: {score_info}, Method: {method_info})")
                                details_run.italic = True
                                details_run.font.size = Pt(9)
                                details_run.font.color.rgb = RGBColor(100, 100, 100) # Gray color

                    doc.add_paragraph()  # Add space between sections

                # Add page break between documents *only if there are multiple* documents being exported
                if len(analysis_results) > 1 : doc.add_page_break()

            except ValueError as ve:
                logger.error(f"Data error processing result for {filename} in export_to_word: {str(ve)}")
                doc.add_heading(f"Error Processing {filename}", 2)
                p = doc.add_paragraph()
                error_run = p.add_run(f"Error processing analysis data: {str(ve)}")
                error_run.font.color.rgb = RGBColor(255, 0, 0)
                p = doc.add_paragraph()
                p.add_run("The system encountered a data error while trying to export this document's analysis.")
                if len(analysis_results) > 1 : doc.add_page_break()

            except Exception as e:
                logger.error(f"Unexpected error processing result for {filename} in export_to_word: {str(e)}", exc_info=True)
                doc.add_heading(f"Unexpected Error Processing {filename}", 2)
                p = doc.add_paragraph()
                error_run = p.add_run(f"Unexpected error processing analysis: {str(e)}")
                error_run.font.color.rgb = RGBColor(255, 0, 0)
                p = doc.add_paragraph()
                p.add_run("The system encountered an unexpected error while trying to export this document's analysis.")
                if len(analysis_results) > 1 : doc.add_page_break()

        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output.getvalue()

    except ImportError:
        logger.error("Export to Word failed: 'python-docx' is not installed.")
        # Create a simple error document indicating the missing dependency
        doc = Document()
        doc.add_heading('Export Error', 0)
        p = doc.add_paragraph()
        error_run = p.add_run("Failed to export analysis results because the 'python-docx' library is not installed.")
        error_run.font.color.rgb = RGBColor(255, 0, 0)
        p.add_run("\nPlease install it using: pip install python-docx")
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    except Exception as e:
        logger.error(f"Error exporting to Word (Outer Try/Except): {str(e)}", exc_info=True)

        # Create a simple generic error document
        doc = Document()
        doc.add_heading('Export Error', 0)
        p = doc.add_paragraph()
        error_run = p.add_run(f"Failed to export analysis results due to an unexpected error: {str(e)}")
        error_run.font.color.rgb = RGBColor(255, 0, 0)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        return output.getvalue()

def process_file_wrapper(args):
    """
    Wrapper for processing a single file: decompose prompt, run RAG & analysis per sub-prompt, aggregate results.
    Now uses precomputed embeddings when available.
    """
    (
        uploaded_file_data,
        filename,
        user_prompt,
        use_advanced_extraction, # Keep if PDFProcessor uses it
        # --- NEW: Pass preprocessed data if available --- 
        preprocessed_data_for_file
    ) = args

    if embedding_model is None:
        logger.error(f"Skipping processing for {filename}: Embedding model not loaded.")
        return {"filename": filename, "error": "Embedding model failed to load.", "annotated_pdf": None, "verification_results": {}, "phrase_locations": {}, "ai_analysis": json.dumps({"error": "Embedding model failed to load."})}

    logger.info(f"Thread {threading.current_thread().name} starting processing for: {filename}")
    start_time = datetime.now()

    try:
        # Check for preprocessed data
        # Use the data passed via args instead of session_state
        preprocessed_data = preprocessed_data_for_file
        
        if preprocessed_data:
            # Use precomputed data
            logger.info(f"Using preprocessed data for {filename} (from {preprocessed_data.get('timestamp')})")
            chunks = preprocessed_data.get("chunks")
            precomputed_embeddings = preprocessed_data.get("chunk_embeddings")
            valid_chunk_indices = preprocessed_data.get("valid_chunk_indices")
            original_pdf_bytes_for_annotation = preprocessed_data.get("original_bytes")
        else:
            # Fall back to processing data on the fly
            logger.warning(f"No preprocessed data found for {filename}, processing on the fly")
            file_extension = Path(filename).suffix.lower()
            processor = None
            original_pdf_bytes_for_annotation = None

            # --- File Handling & Chunk Extraction ---
            if file_extension == ".pdf":
                original_pdf_bytes_for_annotation = uploaded_file_data
                processor = PDFProcessor(uploaded_file_data)
                chunks, _ = processor.extract_structured_text_and_chunks()
            elif file_extension == ".docx":
                word_processor = WordProcessor(uploaded_file_data)
                pdf_bytes = word_processor.convert_to_pdf_bytes()
                if not pdf_bytes: raise ValueError("Failed to convert DOCX to PDF.")
                original_pdf_bytes_for_annotation = pdf_bytes
                processor = PDFProcessor(pdf_bytes) # Process the converted PDF
                chunks, _ = processor.extract_structured_text_and_chunks()
            else:
                # Should not happen due to file uploader limits, but handle defensively
                raise ValueError(f"Unsupported file type: {file_extension}")
                
            # Set precomputed data to None to indicate computation is needed
            precomputed_embeddings = None
            valid_chunk_indices = None

        if not chunks:
            logger.warning(f"No chunks extracted for {filename}, cannot proceed with analysis.")
            b64_pdf = base64.b64encode(original_pdf_bytes_for_annotation).decode() if original_pdf_bytes_for_annotation else None
            return {"filename": filename, "error": "No text chunks could be extracted.", "annotated_pdf": b64_pdf, "verification_results": {}, "phrase_locations": {}, "ai_analysis": json.dumps({"error": "Failed to extract text chunks."})}

        analyzer = DocumentAnalyzer()

        # --- Step 1: Decompose Prompt ---
        decomposed_tasks = []
        try:
            decomposed_tasks = run_async(analyzer.decompose_prompt(user_prompt))
            # Check if fallback occurred (returns list with one item and default title)
            if len(decomposed_tasks) == 1 and decomposed_tasks[0]['title'] == "Overall Analysis":
                 logger.info(f"Decomposition failed or resulted in single prompt. Processing original prompt for {filename}.")
                 # Keep the single item in decomposed_tasks
            else:
                 logger.info(f"Decomposed prompt for {filename} into {len(decomposed_tasks)} sub-prompts with titles.")
        except Exception as decomp_err:
            logger.error(f"Prompt decomposition failed for {filename}: {decomp_err}. Processing original prompt as fallback.", exc_info=True)
            decomposed_tasks = [{"title": "Overall Analysis", "sub_prompt": user_prompt}] # Fallback

        # Change structure: List of (title, sub_prompt, analysis_json_str) tuples
        individual_analysis_results: List[Tuple[str, str, str]] = []
        all_relevant_chunk_ids = set()

        # --- Steps 2 & 3: Loop - Targeted RAG & Focused Analysis per Sub-Prompt ---
        for i, task in enumerate(decomposed_tasks):
             sub_start_time = datetime.now()
             task_title = task['title']
             sub_prompt = task['sub_prompt']
             logger.info(f"Processing sub-task {i+1}/{len(decomposed_tasks)} for {filename}: Title='{task_title}', Prompt='{sub_prompt[:60]}...'")
             
             # Step 2: RAG for this sub-prompt (using precomputed embeddings if available)
             retrieved_rag_chunks = retrieve_relevant_chunks(
                 sub_prompt, chunks, embedding_model, RAG_TOP_K,
                 precomputed_embeddings=precomputed_embeddings,
                 valid_chunk_indices=valid_chunk_indices
             )
             # --- Update how relevant_text and chunk IDs are handled ---
             # Construct relevant_text string for the analysis function
             relevant_text = "\\n\\n---\\n\\n".join(
                 f"[Chunk ID: {chunk.get('chunk_id', 'N/A')}, Page: {chunk.get('page_num', -1) + 1}, Score: {chunk.get('score', 0):.4f}]\\n{chunk.get('text', '')}"
                 for chunk in retrieved_rag_chunks
             )
             # Extract chunk IDs from the retrieved list of dicts
             relevant_chunk_ids = [chunk.get("chunk_id", f"unknown_{i}") for i, chunk in enumerate(retrieved_rag_chunks)]
             all_relevant_chunk_ids.update(relevant_chunk_ids) # Add retrieved chunks to the set

             # Step 3: AI Analysis for this sub-prompt
             analysis_json_str = run_async(
                 analyzer.analyze_document(relevant_text, filename, sub_prompt)
             )
             # Append title, sub_prompt, and result
             individual_analysis_results.append((task_title, sub_prompt, analysis_json_str))
             sub_elapsed = (datetime.now() - sub_start_time).total_seconds()
             logger.info(f"Sub-task {i+1}/{len(decomposed_tasks)} completed in {sub_elapsed:.2f}s")
             
             # If we have a status_container from the main thread, update it with completion


        # --- Step 4: Aggregate Results ---
        logger.info(f"Aggregating {len(individual_analysis_results)} analysis results for {filename}.")
        # If we have a status_container from the main thread, update it
            
        aggregated_ai_analysis_json_str = aggregate_analysis_results(
            individual_analysis_results, filename, user_prompt # Pass the modified list
        )

        # --- Step 5: Verification & Annotation (on aggregated results) ---
        # If we have a status_container from the main thread, update it
                    
        # Create a processor if none exists (if preprocessed data was used)
        if preprocessed_data and 'processor' not in locals():
            processor = PDFProcessor(original_pdf_bytes_for_annotation)
            
        # Verification uses the *original* processor instance with all chunks
        logger.info(f"Verifying phrases from aggregated analysis for {filename}.")
        verification_results, phrase_locations = processor.verify_and_locate_phrases(
            aggregated_ai_analysis_json_str # Use aggregated result
        )

        # Annotation uses the *original* PDF bytes
        logger.info(f"Adding annotations to PDF for {filename}.")
        # Re-initialize PDFProcessor ONLY if original bytes differ from processed bytes (e.g., DOCX conversion)
        # In this structure, original_pdf_bytes_for_annotation holds the correct bytes.
        annotation_processor = PDFProcessor(original_pdf_bytes_for_annotation)
        annotated_pdf_bytes = annotation_processor.add_annotations(phrase_locations)

        end_time = datetime.now()
        total_duration = (end_time - start_time).total_seconds()
        logger.info(f"Successfully processed {filename} in {total_duration:.2f} seconds.")
        

        return {
            "filename": filename,
            "annotated_pdf": base64.b64encode(annotated_pdf_bytes).decode() if annotated_pdf_bytes else None,
            "verification_results": verification_results,
            "phrase_locations": phrase_locations,
            "ai_analysis": aggregated_ai_analysis_json_str, # Return aggregated result
            "retrieved_chunk_ids": list(all_relevant_chunk_ids), # Unique list
        }

    except Exception as e:
        logger.error(f"Error in process_file_wrapper for {filename}: {str(e)}", exc_info=True)
            
        # Try to get original bytes if available for error display
        err_pdf_bytes = None
        if 'original_pdf_bytes_for_annotation' in locals() and original_pdf_bytes_for_annotation:
             err_pdf_bytes = original_pdf_bytes_for_annotation
        elif 'uploaded_file_data' in locals() and file_extension == '.pdf':
             err_pdf_bytes = uploaded_file_data

        b64_err_pdf = base64.b64encode(err_pdf_bytes).decode() if err_pdf_bytes else None
        return {
            "filename": filename,
            "error": f"Processing failed: {str(e)}",
            "annotated_pdf": b64_err_pdf,
            "verification_results": {},
            "phrase_locations": {},
            "ai_analysis": json.dumps({"error": f"Failed to process file: {e}"}),
        }


def display_page():
    """Main function to display the Streamlit page with prompt decomposition."""
    # --- Initialize Session State ---
    if "analysis_results" not in st.session_state: st.session_state.analysis_results = []
    if "show_pdf" not in st.session_state: st.session_state.show_pdf = False
    if "pdf_page" not in st.session_state: st.session_state.pdf_page = 1
    if "pdf_bytes" not in st.session_state: st.session_state.pdf_bytes = None
    if "current_pdf_name" not in st.session_state: st.session_state.current_pdf_name = None
    if "user_prompt" not in st.session_state: st.session_state.user_prompt = ""
    if "use_advanced_extraction" not in st.session_state: st.session_state.use_advanced_extraction = False
    if "last_uploaded_filenames" not in st.session_state: st.session_state.last_uploaded_filenames = set()
    if "uploaded_file_objects" not in st.session_state: st.session_state.uploaded_file_objects = []
    if "preprocessed_data" not in st.session_state: st.session_state.preprocessed_data = {}
    if "preprocessing_status" not in st.session_state: st.session_state.preprocessing_status = {}
    # --- NEW: Initialize chat messages ---    
    if "chat_messages" not in st.session_state: st.session_state.chat_messages = []
    # --- NEW: Flag for user file changes ---
    if "file_selection_changed_by_user" not in st.session_state: st.session_state.file_selection_changed_by_user = False
    # --- NEW: Store files temporarily during change handling ---
    if "current_file_objects_from_change" not in st.session_state: st.session_state.current_file_objects_from_change = None 
    # --- NEW: Flag for auto-scroll ---
    if "results_just_generated" not in st.session_state: st.session_state.results_just_generated = False

    # Check if results exist
    if st.session_state.get("analysis_results"):
        # --- RESULTS VIEW ---
        st.markdown(
            "<p style='text-align: center; font-style: italic;'>AI Powered Document Intelligence</p>",
            unsafe_allow_html=True,
        )

        if st.button("🚀 Start New Analysis", key="new_analysis_button", use_container_width=True, type="primary"):
            # Clear relevant session state variables
            keys_to_clear = [
                "analysis_results", "pdf_bytes", "show_pdf", 
                "current_pdf_name", "chat_messages", "results_just_generated",
                "user_prompt", "uploaded_file_objects", "last_uploaded_filenames",
                "preprocessed_data", "preprocessing_status"
            ]
            for key in keys_to_clear:
                if key in st.session_state:
                    del st.session_state[key]
            logger.info("Cleared state for new analysis.")
            st.rerun()

        # Display Results Section (moved from outside)
        st.divider()

        results_to_display = st.session_state.get("analysis_results", [])
        errors = [r for r in results_to_display if isinstance(r, dict) and "error" in r]
        success_results = [r for r in results_to_display if isinstance(r, dict) and "error" not in r]

        # Status Summary
        total_processed = len(results_to_display)
        if errors:
            if not success_results: st.error(f"Processing failed for all {total_processed} file(s). See details below.")
            else: st.warning(f"Processing complete for {total_processed} file(s). {len(success_results)} succeeded, {len(errors)} failed.")
        # Removed the success message that was here

        # Error Details Expander
        if errors:
            with st.expander("⚠️ Processing Errors", expanded=True):
                for error_res in errors:
                    st.error(f"**{error_res.get('filename', 'Unknown File')}**: {error_res.get('error', 'Unknown error details.')}")

        # Display Successful Analysis
        if success_results:
            display_analysis_results(success_results)
        elif not errors:
            st.warning("Processing finished, but no primary analysis content was generated.")

        # Auto-scroll logic (only runs when results are first shown)
        if st.session_state.get("results_just_generated", False):
            js = """
            <script>
                setTimeout(function() {
                    const anchor = document.getElementById('results-anchor');
                    if (anchor) {
                        console.log("Scrolling to results anchor...");
                        anchor.scrollIntoView({ behavior: 'smooth', block: 'start' });
                    }
                }, 100);
            </script>
            """
            st.components.v1.html(js, height=0)
            st.session_state.results_just_generated = False # Reset flag after scroll

    else:
        # --- INPUT VIEW ---

        st.markdown(
            "<p style='text-align: center; font-style: italic;'>AI Powered Document Intelligence</p>",
            unsafe_allow_html=True,
        )

        # Check if embedding model loaded successfully (important for input view too)
        if embedding_model is None:
            st.error(
                "Embedding model failed to load. Document processing is disabled. "
                "Please check logs and ensure dependencies ('sentence-transformers', 'torch', etc.) are installed correctly."
            )
            return # Stop further UI rendering

        # File Upload Callback
        def handle_file_change():
            current_files = st.session_state.get("file_uploader_decompose", [])
            st.session_state.current_file_objects_from_change = current_files
            st.session_state.file_selection_changed_by_user = True
            logger.debug(f"handle_file_change: Stored {len(current_files) if current_files else 0} files. Flag set.")

        # File Uploader
        uploaded_files = st.file_uploader(
            "Upload PDF or Word files",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key="file_uploader_decompose",
            on_change=handle_file_change,
        )
        
        # Create a placeholder for preprocessing status or features
        preprocessing_or_features_container = st.empty()
        
        # File Change Logic
        if st.session_state.file_selection_changed_by_user:
            logger.debug("Processing detected file change from user action.")
            st.session_state.file_selection_changed_by_user = False
            current_files = st.session_state.current_file_objects_from_change
            current_uploaded_filenames = set(f.name for f in current_files) if current_files else set()
            last_filenames = st.session_state.get('last_uploaded_filenames', set())

            if current_uploaded_filenames != last_filenames:
                logger.info(f"Actual file change confirmed: New={current_uploaded_filenames - last_filenames}, Removed={last_filenames - current_uploaded_filenames}")
                new_files = current_uploaded_filenames - last_filenames
                removed_files = last_filenames - current_uploaded_filenames
                st.session_state.uploaded_file_objects = current_files
                st.session_state.last_uploaded_filenames = current_uploaded_filenames

                for removed_file in removed_files:
                    if removed_file in st.session_state.preprocessed_data:
                        del st.session_state.preprocessed_data[removed_file]
                        if removed_file in st.session_state.preprocessing_status:
                            del st.session_state.preprocessing_status[removed_file]
                        logger.info(f"Removed preprocessing data for {removed_file}")
                
                st.session_state.analysis_results = [] # Clear any old results if files change
                st.session_state.chat_messages = [] # Clear chat too
                logger.info("Cleared relevant state due to file change.")

                if new_files:
                    # Create a status indicator in place of the features container
                    with preprocessing_or_features_container.container():
                        with st.status(f"Preprocessing {len(new_files)} document(s)...", expanded=True) as status:
                            preprocessing_failed = False
                            success_count = 0
                            
                            for i, filename in enumerate(sorted(list(new_files))):
                                file_obj = next((f for f in current_files if f.name == filename), None)
                                if file_obj:
                                    try:
                                        # Update status with current file
                                        status.update(label=f"Preprocessing file {i+1}/{len(new_files)}: {filename}")
                                        st.write(f"Processing {filename}...")
                                        
                                        file_data = file_obj.getvalue()
                                        result = preprocess_file(
                                            file_data, 
                                            filename,
                                            st.session_state.get("use_advanced_extraction", False)
                                        )
                                        st.session_state.preprocessing_status[filename] = result
                                        logger.info(f"Preprocessed {filename}: {result['status']}")
                                        
                                        if result['status'] == 'success':
                                            success_count += 1
                                            st.write(f"✅ {filename} processed successfully.")
                                        elif result['status'] == 'warning':
                                            st.write(f"⚠️ {filename} processed with warnings: {result['message']}")
                                            preprocessing_failed = True
                                        else:
                                            st.write(f"❌ Error processing {filename}: {result['message']}")
                                            preprocessing_failed = True
                                            
                                    except Exception as e:
                                        logger.error(f"Failed to preprocess {filename}: {str(e)}", exc_info=True)
                                        st.session_state.preprocessing_status[filename] = {"status": "error", "message": f"Failed to preprocess: {str(e)}"}
                                        st.write(f"❌ Error processing {filename}: {str(e)}")
                                        preprocessing_failed = True
                            
                            # Update status based on results
                            if preprocessing_failed:
                                if success_count > 0:
                                    status.update(label=f"Preprocessing complete with issues. {success_count}/{len(new_files)} files processed successfully.", state="warning", expanded=False)
                                else:
                                    status.update(label="Preprocessing failed. Please check the errors and try again.", state="error", expanded=False)
                            else:
                                status.update(label=f"Preprocessing complete! {success_count}/{len(new_files)} files processed successfully.", state="complete", expanded=False)
                else:
                    logger.debug("File change flag was True, but filename sets match. Ignoring spurious flag.")
                         
            st.session_state.current_file_objects_from_change = None

        # Welcome Features Section - Only show before files are processed and when not preprocessing
        if not st.session_state.get("preprocessed_data"):            
            # Display features grid only if not currently preprocessing files
            with preprocessing_or_features_container.container():
                st.markdown(f"""           
                <div class="features-container">
                    <div class="features-row">
                        <div class="feature-text">{check_img} Upload your documents, ask questions and get AI analysis with verified responses.</div>
                        <div class="feature-text">{check_img} Get analysis results from multiple documents at once with the same prompt.</div>
                    </div>
                    <div class="features-row">
                        <div class="feature-text">{check_img} Ask Followup questions and talk to your documents to get more details and insights.</div>
                        <div class="feature-text">{check_img} Export annotated documents and query results in Excel or Word format for further analysis.</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

        # Analysis Inputs - Only show if preprocessed data exists
        if st.session_state.get("preprocessed_data"):
            with st.container(border=False):
                st.session_state.user_prompt = st.text_area(
                    "Analysis Prompt",
                    placeholder="Enter your analysis instructions...",
                    height=150,
                    key="prompt_input_decompose",
                    value=st.session_state.get("user_prompt", ""),
                )

            # Process Button
            process_button_disabled = (
                embedding_model is None
                or not st.session_state.get('uploaded_file_objects')
                or not st.session_state.get('user_prompt', '').strip()
            )
            if st.button("Process Documents", type="primary", use_container_width=True, disabled=process_button_disabled):
                files_to_process = st.session_state.get("uploaded_file_objects", [])
                current_user_prompt = st.session_state.get("user_prompt", "")
                current_use_advanced = st.session_state.get("use_advanced_extraction", False)

                if not files_to_process: st.warning("Please upload one or more documents.")
                elif not current_user_prompt.strip(): st.error("Please enter an Analysis Prompt.")
                else:
                    st.session_state.analysis_results = [] # Clear previous results before processing
                    st.session_state.show_pdf = False
                    st.session_state.pdf_bytes = None
                    st.session_state.current_pdf_name = None

                    total_files = len(files_to_process)
                    overall_start_time = datetime.now()
                    results_placeholder = [None] * total_files
                    file_map = {i: f.name for i, f in enumerate(files_to_process)}

                    process_args = []
                    files_read_ok = True
                    for i, uploaded_file in enumerate(files_to_process):
                        try:
                            file_data = uploaded_file.getvalue()
                            # Add the preprocessed data for this file to the args
                            preprocessed_file_data = st.session_state.get("preprocessed_data", {}).get(uploaded_file.name)
                            process_args.append(
                                (file_data, uploaded_file.name, current_user_prompt, current_use_advanced, preprocessed_file_data)
                            )
                        except Exception as read_err:
                            logger.error(f"Failed to read file {uploaded_file.name}: {read_err}", exc_info=True)
                            st.error(f"Failed to read file {uploaded_file.name}. Please re-upload.")
                            results_placeholder[i] = {"filename": uploaded_file.name, "error": f"Failed to read file: {read_err}"}
                            files_read_ok = False

                    if files_read_ok and process_args:
                        files_to_run_count = len(process_args)
                        
                        # Create expander for document processing status
                        with st.expander("Document Processing Status", expanded=True):
                            status_container = st.empty()
                            status_container.info(f"Starting to process {files_to_run_count} document(s)...")
                            
                            with st.spinner("Analysing Query...", show_time=True):
                                processed_indices = set()
                                def run_task_with_index(item_index: int, args_tuple: tuple):
                                    filename = args_tuple[1]
                                    logger.info(f"Thread {threading.current_thread().name} starting task for index {item_index} ({filename})")
                                    try:
                                        result = process_file_wrapper(args_tuple)
                                        logger.info(f"Thread {threading.current_thread().name} finished task for index {item_index} ({filename})")
                                        return item_index, result
                                    except Exception as thread_err:
                                        logger.error(f"Unhandled error in thread task for index {item_index} ({filename}): {thread_err}", exc_info=True)
                                        return item_index, {"filename": filename, "error": f"Unhandled thread error: {thread_err}"}

                                try:
                                    if ENABLE_PARALLEL and len(process_args) > 1:
                                        logger.info(f"Executing {len(process_args)} tasks in parallel with max workers: {MAX_WORKERS}")
                                        # Process each file in parallel
                                        with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                                            future_to_index = {executor.submit(run_task_with_index, i, args): i for i, args in enumerate(process_args)}
                                            
                                            for future in concurrent.futures.as_completed(future_to_index):
                                                original_index = future_to_index[future]
                                                processed_indices.add(original_index)
                                                fname = file_map.get(original_index, f"File at index {original_index}")
                                                try:
                                                    _, result_data = future.result()
                                                    results_placeholder[original_index] = result_data
                                                except Exception as exc:
                                                    logger.error(f'Task for index {original_index} ({fname}) failed: {exc}', exc_info=True)
                                                    results_placeholder[original_index] = {"filename": fname, "error": f"Task execution failed: {exc}"}
                                    else:
                                        logger.info(f"Processing {files_to_run_count} task(s) sequentially.")
                                        for i, arg_tuple in enumerate(process_args):
                                            original_index = i
                                            processed_indices.add(original_index)
                                            try:
                                                _, result_data = run_task_with_index(original_index, arg_tuple)
                                                results_placeholder[original_index] = result_data
                                            except Exception as seq_exc:
                                                 fname = file_map.get(original_index, f"File at index {original_index}")
                                                 logger.error(f'Sequential task for index {original_index} ({fname}) failed: {seq_exc}', exc_info=True)
                                                 results_placeholder[original_index] = {"filename": fname, "error": f"Task execution failed: {seq_exc}"}

                                except Exception as pool_err:
                                     logger.error(f"Error during task execution setup/management: {pool_err}", exc_info=True)
                                     st.error(f"Error during processing: {pool_err}. Some files may not have been processed.")
                                     for i in range(total_files):
                                          if i not in processed_indices and results_placeholder[i] is None:
                                               fname = file_map.get(i, f"File at index {i}")
                                               results_placeholder[i] = {"filename": fname, "error": f"Processing cancelled due to execution error: {pool_err}"}

                                final_results = [r for r in results_placeholder if r is not None]
                                st.session_state.analysis_results = final_results
                                total_time = (datetime.now() - overall_start_time).total_seconds()
                                success_count = len([r for r in final_results if isinstance(r, dict) and "error" not in r])
                                logger.info(f"Processing batch complete. Processed {success_count}/{total_files} files successfully in {total_time:.2f}s.")
                                status_container.success(f"Processing complete! Processed {success_count}/{total_files} files successfully in {total_time:.2f} seconds.")

                        # Outside the expander, handle post-processing like PDF loading
                        first_success = next((r for r in final_results if isinstance(r, dict) and "error" not in r), None)
                        if first_success and first_success.get("annotated_pdf"):
                            try:
                                pdf_bytes = base64.b64decode(first_success["annotated_pdf"])
                                update_pdf_view(pdf_bytes=pdf_bytes, page_num=1, filename=first_success.get("filename"))
                            except Exception as decode_err:
                                logger.error(f"Failed to decode/set initial PDF: {decode_err}", exc_info=True)
                                st.error("Failed to load initial PDF view.")
                                st.session_state.show_pdf = False
                        elif first_success:
                             logger.warning("First successful result missing annotated PDF data.")
                             st.warning("Processing complete, but couldn't display the first annotated document.")
                             st.session_state.show_pdf = False
                        else:
                             logger.warning("No successful results found. No initial PDF view shown.")
                             st.session_state.show_pdf = False

                    # Set flag to trigger scroll on next rerun
                    if success_count > 0:
                        st.session_state.results_just_generated = True

                    # Rerun to display results / updated PDF view state
                    st.rerun()


# --- Main Execution Guard ---
if __name__ == "__main__":
    # Check model again before displaying page, though initial check should handle most cases
    if embedding_model is not None:
        display_page()
    else:
        # If the model failed, display_page() will show an error,
        # but we can add a log here just in case.
        logger.critical("Application cannot start because the embedding model failed to load.")
        # Avoid st calls here if it might not be initialized